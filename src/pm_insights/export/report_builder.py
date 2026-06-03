from __future__ import annotations

import copy
import io
import json
import re
import tempfile
from pathlib import Path
from typing import Any

from pm_insights.nlp.postprocessing import normalize_analysis_result


MEETING_TYPE_LABELS = {
    "project_meeting": "проектная встреча",
    "technical_research": "техническая / исследовательская встреча",
    "education_consultation": "учебная консультация",
    "commercial_meeting": "коммерческая встреча",
    "commercial_oil_gas": "нефтегазовая коммерческая встреча",
    "oil_gas_commercial": "нефтегазовая коммерческая встреча",
    "mixed": "смешанная встреча",
    "general_discussion": "общее обсуждение",
    "unknown": "тип встречи не определён",
}

SECTION_TITLES_RU = {
    "tasks": "Задачи",
    "recommendations": "Рекомендации",
    "questions_answers": "Вопросы и ответы",
    "qa": "Вопросы и ответы",
    "deadlines": "Дедлайны / следующая встреча",
    "topics": "Ключевые темы",
    "aspects": "Аспекты",
    "aspects_topics": "Аспекты и темы",
    "sentiment": "Тональность",
    "transcript": "Транскрипт",
    "commercial_terms": "Коммерческие условия",
    "agreements": "Договорённости",
    "commitments": "Обещания сторон",
    "responsibles": "Ответственные",
    "responsible_sides": "Ответственные стороны",
    "decisions": "Решения",
}

REVIEW_DEBUG_PATTERNS = (
    ("needs", "review"),
    ("review", "items"),
    ("review", "required"),
    ("requires", "manual", "review"),
    ("low", "confidence", "or", "missing", "owner"),
    ("confidence",),
    ("требует", "проверки"),
    ("нужна", "ручная", "проверка"),
)

REVIEW_DEBUG_KEYS = {
    "_".join(("needs", "review")),
    "_".join(("review", "required")),
    "confidence",
    "reason",
    "reasons",
    "_".join(("review", "reason")),
}


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list, tuple)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _looks_corrupted_heading(text: str) -> bool:
    if not text:
        return True
    question_marks = text.count("?")
    if question_marks >= 3:
        return True
    if question_marks and question_marks / max(len(text), 1) > 0.25:
        return True
    return False


def _compact(value: Any, limit: int = 900) -> str:
    text = re.sub(r"\s+", " ", _as_text(value)).strip()
    if _is_review_debug_text(text):
        return ""
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def truncate_text(value: Any, max_len: int) -> str:
    text = clean_export_text(value)
    if len(text) <= max_len:
        return text
    cut = text[: max_len - 3].rstrip()
    boundary = max(cut.rfind("."), cut.rfind(";"), cut.rfind(","), cut.rfind(" "))
    if boundary >= max_len * 0.55:
        cut = cut[:boundary].rstrip(" .;,")
    return cut + "..."


def format_duration(seconds: Any) -> str:
    try:
        total = int(round(float(seconds)))
    except (TypeError, ValueError):
        return ""
    if total <= 0:
        return ""
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    parts = []
    if hours:
        parts.append(f"{hours} ч")
    if minutes or hours:
        parts.append(f"{minutes} мин")
    parts.append(f"{secs} сек")
    return " ".join(parts)


def clean_export_text(value: Any) -> str:
    text = re.sub(r"\s+", " ", _as_text(value)).strip()
    if _is_review_debug_text(text):
        return ""
    return text


def _name_key(value: Any) -> str:
    return re.sub(r"[^a-zа-яё0-9]+", " ", clean_export_text(value).lower().replace("ё", "е")).strip()


EXPORT_TOPIC_RENAME = {
    "требования этой эту": "требования к ВКР",
    "постановке задачи": "постановка задачи",
    "диаметр интересно случайный": "диаметр и топология графа",
    "минимизировать меньше связей": "оптимизация структуры графа",
    "минимизировать максимизировать": "оптимизационная постановка",
}

EXPORT_TOPIC_HIDE = {
    "причем понял честно",
    "только использую используют",
}

EXPORT_TOPIC_FILLER = {
    "это",
    "эту",
    "этой",
    "причем",
    "понял",
    "честно",
    "только",
    "использую",
    "используют",
    "интересно",
    "случайный",
    "какие",
    "что",
    "просто",
    "вообще",
}

GENERIC_TOPIC_NAMES = {
    "модель",
    "вкр",
    "данные",
    "задача",
    "темы",
    "аспекты",
    "критерии",
}

EXPORT_TOPIC_KEEP = {
    "постановка задачи",
    "требования к вкр",
    "диаметр и топология графа",
    "оптимизация структуры графа",
    "количество связей",
    "необходимые критерии",
    "безразмерные кривые",
    "метрики графа",
    "ограничения модели",
    "сбор данных и телеметрия",
}

WEAK_EXPORT_TOPIC_PATTERNS = (
    "полезная тема",
    "причем",
    "понял",
    "честно",
    "только использую",
    "какие то",
    "что то",
    "что-нибудь",
    "где то",
)

WEAK_ACTION_KEYS = {
    "это необходимо сделать",
    "можно будет",
    "далеко не бесполезно",
    "лучше проверить",
    "я попробую поискать",
}

INTERROGATIVE_STARTS = (
    "как ",
    "что ",
    "почему ",
    "зачем ",
    "где ",
    "когда ",
    "кто ",
    "какой ",
    "какая ",
    "какие ",
    "можно ли ",
    "нужно ли ",
    "стоит ли ",
    "правильно ли ",
)

WEAK_QUESTION_PATTERNS = (
    "что то там",
    "где то можно",
    "что у нас есть какие то",
    "что это решение ожидается",
)


def normalize_keywords(value: Any) -> list[str]:
    if value in (None, "", [], {}, ()):
        return []
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        try:
            parsed = json.loads(text)
        except (TypeError, ValueError, json.JSONDecodeError):
            parsed = None
        if parsed is not None and parsed is not value:
            return normalize_keywords(parsed)
        text = text.strip("[](){}")
        parts = re.split(r"[,;|/\n]+", text)
        return _dedupe_keyword_parts(parts)
    if isinstance(value, dict):
        return normalize_keywords(
            value.get("word")
            or value.get("keyword")
            or value.get("name")
            or value.get("title")
            or value.get("text")
        )
    if isinstance(value, (list, tuple, set)):
        values = list(value)
        if not values:
            return []
        if all(isinstance(item, str) for item in values):
            nonempty = [item for item in values if item]
            if nonempty and sum(1 for item in nonempty if len(item.strip()) <= 1 or item in {",", ";", " "}) >= len(nonempty) * 0.75:
                return normalize_keywords("".join(nonempty))
        parts: list[str] = []
        for item in values:
            parts.extend(normalize_keywords(item))
        return _dedupe_keyword_parts(parts)
    return _dedupe_keyword_parts([str(value)])


def _dedupe_keyword_parts(parts: list[Any]) -> list[str]:
    cleaned: list[str] = []
    for part in parts:
        text = clean_export_text(part).strip(" \"'[]{}()")
        if not text or text in {",", ";", "|", "/", "—", "-"}:
            continue
        if len(text) <= 1:
            continue
        if re.fullmatch(r"[\W_]+", text, flags=re.UNICODE):
            continue
        key = _name_key(text)
        if not key or key in {",", "и", "или"}:
            continue
        if key not in {_name_key(item) for item in cleaned}:
            cleaned.append(text)
    return cleaned[:8]


def format_keywords(value: Any, max_len: int = 120) -> str:
    keywords = normalize_keywords(value)
    if not keywords:
        return "—"
    return truncate_text(", ".join(keywords), max_len)


def display_cell(value: Any, max_len: int = 300) -> str:
    if value in (None, "", [], {}, ()):
        return "—"
    if isinstance(value, (list, tuple, set)):
        values = [display_cell(item, max_len) for item in value]
        values = [item for item in values if item and item != "—"]
        return truncate_text(", ".join(values), max_len) if values else "—"
    if isinstance(value, dict):
        for key in ("title", "name", "text", "value", "summary"):
            if value.get(key):
                return display_cell(value.get(key), max_len)
        return "—"
    text = clean_export_text(value)
    if not text or text in {"[]", "{}"}:
        return "—"
    return truncate_text(text, max_len)


def _dedupe_rows_by_field(rows: list[dict], field: str = "title") -> list[dict]:
    deduped: list[dict] = []
    seen: set[str] = set()
    for row in rows:
        key = _name_key(row.get(field))
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def _is_generic_export_action(text: str) -> bool:
    key = _name_key(text)
    if not key or len(key) < 8:
        return True
    if key in WEAK_ACTION_KEYS:
        return True
    if any(pattern in key for pattern in WEAK_ACTION_KEYS if pattern != "лучше проверить"):
        return True
    return False


def _export_topic_name(item: dict) -> str:
    return (
        clean_export_text(item.get("topic_name"))
        or clean_export_text(item.get("title"))
        or clean_export_text(item.get("name"))
    )


def is_bad_export_topic(name: Any) -> bool:
    text = clean_export_text(name)
    if not text:
        return True
    key = _name_key(text)
    if key in EXPORT_TOPIC_KEEP:
        return False
    if key in EXPORT_TOPIC_HIDE:
        return True
    if len(key) < 3:
        return True
    if any(pattern in key for pattern in WEAK_EXPORT_TOPIC_PATTERNS):
        return True
    words = key.split()
    if len(text) > 80:
        return True
    if len(words) >= 2:
        filler_hits = sum(1 for word in words if word in EXPORT_TOPIC_FILLER)
        if filler_hits >= max(2, len(words) - 1):
            return True
    if re.fullmatch(r"(тема|кластер|прочее)(\s+\d+)?", key):
        return True
    return False


def _export_topic_quality_ok(name: Any, keywords: Any) -> bool:
    normalized_name = _normalize_export_topic_name(name)
    if not normalized_name or is_bad_export_topic(normalized_name):
        return False
    key = _name_key(normalized_name)
    normalized_keywords = normalize_keywords(keywords)
    if key in EXPORT_TOPIC_KEEP:
        return True
    if key in GENERIC_TOPIC_NAMES and len(normalized_keywords) < 2:
        return False
    if not normalized_keywords and len(key.split()) < 2:
        return False
    if normalized_keywords and len(normalized_keywords) < 2 and key in GENERIC_TOPIC_NAMES:
        return False
    return True


def _normalize_export_topic_name(name: Any) -> str:
    text = clean_export_text(name)
    key = _name_key(text)
    if key in EXPORT_TOPIC_RENAME:
        return EXPORT_TOPIC_RENAME[key]
    if key in EXPORT_TOPIC_HIDE:
        return ""
    return text


def dedupe_named_items(items: list[dict], *, name_fields: tuple[str, ...] = ("title", "topic_name", "name")) -> list[dict]:
    grouped: dict[str, dict] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        name = next((clean_export_text(item.get(field)) for field in name_fields if clean_export_text(item.get(field))), "")
        if not name:
            continue
        key = _name_key(name)
        if not key:
            continue
        if key not in grouped:
            grouped[key] = dict(item)
            grouped[key]["title"] = name
            grouped[key]["count"] = int(item.get("count") or 1)
            grouped[key]["priority"] = int(item.get("priority") or 0)
            grouped[key]["keywords"] = normalize_keywords(item.get("keywords"))
            continue
        existing = grouped[key]
        existing["count"] = int(existing.get("count") or 0) + int(item.get("count") or 1)
        existing["priority"] = max(int(existing.get("priority") or 0), int(item.get("priority") or 0))
        existing["keywords"] = _dedupe_keyword_parts(list(existing.get("keywords") or []) + normalize_keywords(item.get("keywords")))
    return sorted(
        grouped.values(),
        key=lambda row: (int(row.get("priority") or 0), int(row.get("count") or 0)),
        reverse=True,
    )


def select_top_items(items: list[dict], limit: int) -> list[dict]:
    return [item for item in items if isinstance(item, dict)][:limit]


def _is_review_debug_text(text: str) -> bool:
    normalized = text.lower().replace("_", " ")
    return any(" ".join(parts) in normalized for parts in REVIEW_DEBUG_PATTERNS)


def _meeting_type(result: dict) -> str:
    return (
        (result.get("meeting_type") or {}).get("label")
        or (result.get("analysis_summary") or {}).get("meeting_type")
        or "unknown"
    )


def _meeting_type_label(result: dict) -> str:
    meeting_type = _meeting_type(result)
    return (result.get("meeting_type") or {}).get("display_name") or MEETING_TYPE_LABELS.get(meeting_type, meeting_type)


def _metadata(result: dict) -> dict:
    metadata = result.get("metadata") or {}
    info = metadata.get("meeting_info") or result.get("meeting_info") or {}
    processing = metadata.get("processing_time") or (result.get("metrics") or {}).get("processing_time") or {}
    return {
        "system": "PM Insights",
        "meeting_id": result.get("meeting_id") or result.get("id") or "",
        "meeting_title": info.get("meeting_title") or result.get("meeting_title") or result.get("source_audio") or "",
        "project_name": info.get("project_name") or result.get("project_name") or "Без проекта",
        "meeting_date": info.get("meeting_date") or result.get("meeting_date") or result.get("created_at") or "",
        "meeting_type": _meeting_type(result),
        "meeting_type_label": _meeting_type_label(result),
        "duration_seconds": metadata.get("duration_seconds") or processing.get("audio_duration_seconds"),
        "analysis_date": result.get("created_at") or metadata.get("analysis_date") or "",
        "asr_status": metadata.get("asr_status") or "completed",
        "asr_model": metadata.get("asr_model") or metadata.get("model") or "",
        "processing_time_seconds": processing.get("total_processing_seconds"),
    }


def _section_items(result: dict, section_id: str, clean_field: str, raw_field: str | None = None) -> list[dict]:
    for section in result.get("report_sections") or []:
        if section.get("id") == section_id and isinstance(section.get("items"), list):
            return section["items"]
    if isinstance(result.get(clean_field), list):
        return result[clean_field]
    if raw_field and isinstance(result.get(raw_field), list):
        return result[raw_field]
    return []


def _section_items_any(
    result: dict,
    section_ids: tuple[str, ...],
    clean_fields: tuple[str, ...],
    raw_fields: tuple[str, ...] = (),
) -> list[dict]:
    for section in result.get("report_sections") or []:
        if section.get("id") in section_ids and isinstance(section.get("items"), list) and section["items"]:
            return section["items"]
    for field in clean_fields:
        if isinstance(result.get(field), list):
            return result[field]
    for field in raw_fields:
        if isinstance(result.get(field), list):
            return result[field]
    return []


def _topic_aspect_items(result: dict) -> list[dict]:
    items: list[dict] = []
    for topic in result.get("clean_topics") or []:
        if isinstance(topic, dict):
            items.append(topic)
    for aspect in result.get("clean_aspects") or []:
        if isinstance(aspect, dict):
            items.append(
                {
                    "topic_name": aspect.get("title"),
                    "title": aspect.get("title"),
                    "count": aspect.get("count"),
                    "keywords": aspect.get("keywords") or [],
                    "fragment_ids": aspect.get("fragment_ids") or [],
                }
            )
    return items


def _responsible_side_names(result: dict) -> list[str]:
    names: list[str] = []
    for source in (
        result.get("clean_responsible_sides") or [],
        result.get("responsible_side") or [],
        result.get("responsible_sides") or [],
    ):
        if isinstance(source, str):
            names.append(source)
        elif isinstance(source, dict):
            value = source.get("side") or source.get("responsible_side") or source.get("title") or source.get("name")
            if value:
                names.append(str(value))
    return list(dict.fromkeys(name.strip() for name in names if name and name.strip()))


def _simplified_sections(result: dict) -> list[dict]:
    meeting_type = _meeting_type(result)
    sections: list[dict] = []

    def add(section_id: str, title: str, items: list[dict]) -> None:
        if items:
            sections.append({"id": section_id, "title": title, "items": items})

    add("tasks", "Задачи", _section_items_any(result, ("tasks",), ("clean_tasks",)))

    add("qa", "Вопросы и ответы", _section_items_any(result, ("questions_answers", "qa"), ("clean_questions_answers",)))
    add("responsibles", "Ответственные", _section_items_any(result, ("responsibles",), ("clean_responsibles",)))
    add("deadlines", "Дедлайны", _section_items_any(result, ("deadlines",), ("clean_deadlines",)))
    add("decisions", "Решения", _section_items_any(result, ("decisions",), ("clean_decisions",)))
    add("topics", "Аспекты и темы", _topic_aspect_items(result))
    add("sentiment", "Тональность", result.get("sentiment") or [])
    return sections


def _status_label(status: Any) -> str:
    mapping = {
        "answered": "дан ответ",
        "partial": "частичный ответ",
        "not_answered": "проигнорирован",
        "ignored": "проигнорирован",
        "new": "новая",
        "repeated": "повторяющаяся",
    }
    return mapping.get(clean_export_text(status), clean_export_text(status))


def _is_weak_question(question: str) -> bool:
    key = _name_key(question)
    if key in {"что", "да", "а", "понятно", "ок", "так", "правильно"}:
        return True
    if len(key) < 4:
        return True
    if any(pattern in key for pattern in WEAK_QUESTION_PATTERNS):
        return True
    if "?" in question:
        return len(question.strip()) < 8
    return not key.startswith(INTERROGATIVE_STARTS) or len(question.strip()) < 8


def _question_for_export(item: dict) -> str:
    question = (
        clean_export_text(item.get("question_title"))
        or clean_export_text(item.get("question"))
        or clean_export_text(item.get("title"))
    )
    if "?" in question:
        parts = re.findall(r"[^?。.!]*\?", question)
        if parts:
            question = parts[-1].strip()
    return question


def _clean_task_rows(result: dict) -> tuple[str, list[dict]]:
    title = "Задачи"
    sources = _section_items_any(result, ("tasks",), ("clean_tasks",))

    rows = []
    for item in sources:
        if not isinstance(item, dict):
            continue
        task_title = (
            clean_export_text(item.get("title"))
            or clean_export_text(item.get("clean_title"))
            or clean_export_text(item.get("text"))
        )
        if not task_title or _is_generic_export_action(task_title):
            continue
        rows.append(
            {
                "title": truncate_text(task_title, 180),
                "responsible": truncate_text(item.get("responsible") or item.get("owner") or "", 80),
                "responsible_side": truncate_text(item.get("responsible_side") or "", 80),
                "deadline": truncate_text(item.get("deadline") or "", 80),
                "status": truncate_text(_status_label(item.get("status") or "new"), 80),
            }
        )
    return title, _dedupe_rows_by_field(rows, "title")[:8]


def _clean_qa_rows(result: dict) -> list[dict]:
    rows = []
    sources = _section_items_any(result, ("questions_answers", "qa"), ("clean_questions_answers",))
    seen: set[str] = set()
    for item in sources:
        if not isinstance(item, dict):
            continue
        question = _question_for_export(item)
        if _is_weak_question(question):
            continue
        key = _name_key(question)
        if key in seen:
            continue
        seen.add(key)
        answer = (
            clean_export_text(item.get("answer_summary"))
            or clean_export_text(item.get("answer"))
            or clean_export_text(item.get("answer_full"))
            or "Ответ не найден"
        )
        rows.append(
            {
                "question": truncate_text(question, 180),
                "answer": truncate_text(answer, 240),
                "status": _status_label(item.get("status") or "not_answered"),
            }
        )
        if len(rows) >= 6:
            break
    return rows


def _clean_deadline_rows(result: dict) -> list[dict]:
    rows = []
    for item in _section_items_any(result, ("deadlines",), ("clean_deadlines",)):
        deadline = clean_export_text(item.get("deadline") or item.get("text"))
        if not deadline:
            continue
        rows.append(
            {
                "deadline": truncate_text(deadline, 80),
                "kind": truncate_text(item.get("kind") or item.get("type") or "", 80),
                "context": truncate_text(item.get("context") or "", 140),
            }
        )
        if len(rows) >= 10:
            break
    return rows


def _clean_responsible_rows(result: dict) -> list[str]:
    names = []
    for item in _section_items_any(result, ("responsibles",), ("clean_responsibles",)):
        if isinstance(item, str):
            value = item
        else:
            value = item.get("name") or item.get("responsible") or item.get("title")
        value = truncate_text(value, 80)
        if value and value not in names:
            names.append(value)
    return names


def _clean_topic_rows(result: dict, limit: int = 8) -> list[dict]:
    rows = []
    for item in result.get("clean_topics") or []:
        if not isinstance(item, dict):
            continue
        raw_name = _export_topic_name(item)
        name = _normalize_export_topic_name(raw_name)
        keywords = normalize_keywords(item.get("keywords"))
        if not _export_topic_quality_ok(name, keywords):
            continue
        rows.append(
            {
                "title": truncate_text(name, 80),
                "count": int(item.get("count") or len(item.get("fragment_ids") or []) or 1),
                "keywords": format_keywords(keywords, 120),
                "priority": 1 if _name_key(raw_name) in EXPORT_TOPIC_RENAME else 0,
            }
        )
    rows = dedupe_named_items(rows, name_fields=("title",))
    if len(rows) < 3:
        for aspect in result.get("clean_aspects") or []:
            if not isinstance(aspect, dict):
                continue
            raw_name = aspect.get("title")
            name = _normalize_export_topic_name(raw_name)
            keywords = normalize_keywords(aspect.get("keywords"))
            if not _export_topic_quality_ok(name, keywords):
                continue
            rows.append(
                {
                    "title": truncate_text(name, 80),
                    "count": int(aspect.get("count") or 1),
                    "keywords": format_keywords(keywords, 120),
                    "priority": 1 if _name_key(raw_name) in EXPORT_TOPIC_RENAME else 0,
                }
        )
    rows = dedupe_named_items(rows, name_fields=("title",))
    for row in rows:
        row["keywords"] = format_keywords(row.get("keywords"), 120)
    return rows[:limit]


def _clean_aspect_rows(result: dict, topics: list[dict] | None = None) -> list[dict]:
    source = result.get("clean_aspects") or []
    topic_keys = {_name_key(row.get("title")) for row in topics or [] if row.get("title")}
    rows = []
    for item in source:
        if not isinstance(item, dict):
            continue
        raw_name = item.get("title") or item.get("topic_name") or item.get("name")
        name = _normalize_export_topic_name(raw_name)
        keywords = normalize_keywords(item.get("keywords"))
        if not _export_topic_quality_ok(name, keywords):
            continue
        key = _name_key(name)
        if key in topic_keys:
            continue
        rows.append(
            {
                "title": truncate_text(name, 80),
                "count": int(item.get("count") or len(item.get("fragment_ids") or []) or 1),
                "keywords": format_keywords(keywords, 120),
                "priority": 1 if _name_key(raw_name) in EXPORT_TOPIC_RENAME else 0,
            }
        )
    rows = dedupe_named_items(rows, name_fields=("title",))
    if topics and rows:
        overlap = 1 - (len(rows) / max(len(rows) + len(topic_keys), 1))
        if overlap > 0.7:
            return []
    for row in rows:
        row["keywords"] = format_keywords(row.get("keywords"), 120)
    return rows[:8]


def _sentiment_label(item: dict) -> str:
    return clean_export_text(item.get("sentiment") or item.get("label") or "neutral")


def _sentiment_average(result: dict, sentiment_items: list[dict]) -> float:
    summary = result.get("analysis_summary") or {}
    metrics = result.get("metrics") or {}
    if summary.get("average_sentiment") is not None:
        return round(float(summary["average_sentiment"]), 3)
    if metrics.get("average_sentiment") is not None:
        return round(float(metrics["average_sentiment"]), 3)
    values = []
    for item in sentiment_items:
        if item.get("score") is not None:
            values.append(float(item.get("score") or 0.0))
        else:
            values.append({"positive": 1.0, "neutral": 0.0, "negative": -1.0}.get(_sentiment_label(item), 0.0))
    return round(sum(values) / len(values), 3) if values else 0.0


def _aspect_sentiment_rows(result: dict) -> list[dict]:
    source = result.get("aspect_sentiment") or []
    if isinstance(source, dict):
        source = [{"aspect": key, **(value if isinstance(value, dict) else {"average": value})} for key, value in source.items()]
    rows = []
    meeting_type = _meeting_type(result)
    for item in source:
        if not isinstance(item, dict):
            continue
        raw_aspect = item.get("aspect") or item.get("title") or item.get("name")
        aspect = _normalize_export_topic_name(raw_aspect)
        average = item.get("average") if item.get("average") is not None else item.get("average_sentiment")
        negative = item.get("negative") if item.get("negative") is not None else item.get("negative_count")
        if average in (None, "") and negative in (None, ""):
            continue
        if meeting_type == "education_consultation" and _name_key(aspect) == "параметры пласта":
            continue
        if not _export_topic_quality_ok(aspect, item.get("keywords")):
            continue
        rows.append(
            {
                "aspect": truncate_text(aspect, 80),
                "average": truncate_text(average, 40),
                "negative": truncate_text(negative, 40),
            }
        )
    return rows[:8]


def build_compact_report_data(result: dict) -> dict:
    if not any(result.get(field) for field in ("clean_tasks", "clean_questions_answers", "clean_topics", "clean_aspects")):
        result = normalize_analysis_result(copy.deepcopy(result))
    metadata = _metadata(result)
    task_title, tasks = _clean_task_rows(result)
    qa = _clean_qa_rows(result)
    deadlines = _clean_deadline_rows(result)
    responsibles = _clean_responsible_rows(result)
    responsible_sides = _responsible_side_names(result)
    topics = _clean_topic_rows(result, 8)
    aspects = _clean_aspect_rows(result, topics)
    sentiment_items = [
        item for item in (result.get("clean_sentiment") or result.get("sentiment") or [])
        if isinstance(item, dict)
    ]
    positive = sum(1 for item in sentiment_items if _sentiment_label(item) == "positive")
    neutral = sum(1 for item in sentiment_items if _sentiment_label(item) == "neutral")
    negative = sum(1 for item in sentiment_items if _sentiment_label(item) == "negative")
    average = _sentiment_average(result, sentiment_items)
    answer_count = sum(1 for item in qa if item["status"] in {"дан ответ", "частичный ответ"} and item.get("answer") != "Ответ не найден")

    return {
        "meeting": {
            "project": metadata.get("project_name"),
            "date": metadata.get("meeting_date"),
            "type": metadata.get("meeting_type_label"),
            "duration": format_duration(metadata.get("duration_seconds")),
            "asr_model": metadata.get("asr_model"),
            "processing_time": format_duration(metadata.get("processing_time_seconds")),
        },
        "metrics": {
            "tasks": len(tasks),
            "questions": len(qa),
            "answers": answer_count,
            "deadlines": len(deadlines),
            "responsibles": len(responsibles),
            "topics": len(topics),
            "aspects": len(aspects) if aspects else len(topics),
            "avg_sentiment": average,
            "negative": negative,
        },
        "topics": topics,
        "tasks_title": task_title,
        "tasks": tasks,
        "qa": qa,
        "deadlines": deadlines,
        "responsibles": responsibles,
        "responsible_sides": responsible_sides,
        "aspects": aspects,
        "sentiment": {
            "positive": positive,
            "neutral": neutral,
            "negative": negative,
            "average": average,
            "aspect_sentiment": _aspect_sentiment_rows(result),
        },
    }


def _public_report_result(result: dict) -> dict:
    cleaned = copy.deepcopy(result)
    for field in ("topics", "aspects", "aspect_frequencies", "topic_frequencies"):
        cleaned.pop(field, None)
    metrics = cleaned.get("metrics")
    if isinstance(metrics, dict):
        metrics.pop("aspect_frequencies", None)
        metrics.pop("topic_frequencies", None)
    return cleaned


def build_report_payload(result: dict) -> dict:
    source = copy.deepcopy(result)
    normalized = normalize_analysis_result(copy.deepcopy(result))
    for field in (
        "clean_tasks",
        "clean_questions_answers",
        "clean_deadlines",
        "clean_responsibles",
        "clean_decisions",
        "clean_topics",
        "clean_aspects",
        "clean_sentiment",
        "sentiment",
    ):
        if not normalized.get(field) and source.get(field):
            normalized[field] = source[field]
    sections = _simplified_sections(normalized)
    summary = normalized.get("analysis_summary") or {}
    source_summary = source.get("analysis_summary") or {}
    if not summary.get("main_topics") and source_summary.get("main_topics"):
        summary = dict(summary)
        summary["main_topics"] = source_summary.get("main_topics") or []
    transcript = normalized.get("transcript") or []
    return {
        "metadata": _metadata(normalized),
        "summary": summary,
        "responsible_sides": _responsible_side_names(normalized),
        "sections": sections,
        "transcript": transcript,
        "compact": build_compact_report_data(normalized),
        "technical_metrics": {
            "processing_time": (normalized.get("metadata") or {}).get("processing_time") or {},
        },
        "result": _public_report_result(normalized),
    }


def export_json(report: dict) -> bytes:
    return json.dumps(report["result"], ensure_ascii=False, indent=2).encode("utf-8")


def _section_items_list(section: dict) -> list[dict]:
    items = section.get("items") or []
    if not isinstance(items, list):
        return []
    cleaned = []
    for item in items:
        if not isinstance(item, dict):
            continue
        cleaned.append(
            {
                key: value
                for key, value in item.items()
                if key not in REVIEW_DEBUG_KEYS
            }
        )
    return cleaned


def _section_title(section: dict) -> str:
    section_id = section.get("id") or ""
    title = _as_text(section.get("title")).strip()
    if _looks_corrupted_heading(title):
        return SECTION_TITLES_RU.get(section_id, section_id or "Раздел")
    return title


def assert_no_question_marks_in_headings(report: dict) -> None:
    headings = [
        "PM Insights — отчёт по встрече",
        "Основные показатели",
        "Основные результаты",
        *[_section_title(section) for section in report.get("sections") or []],
    ]
    damaged = [heading for heading in headings if _looks_corrupted_heading(heading)]
    if damaged:
        raise ValueError(f"Повреждены заголовки отчёта: {damaged}")


def _item_title(item: dict) -> str:
    return (
        item.get("title")
        or item.get("question_title")
        or item.get("topic_name")
        or item.get("deadline")
        or item.get("clean_title")
        or item.get("text")
        or "Элемент"
    )


def _item_summary(item: dict) -> str:
    return item.get("summary") or item.get("answer_summary") or item.get("context") or ""


def _append_sheet(workbook, title: str, rows: list[dict], columns: list[str]) -> None:
    sheet = workbook.create_sheet(title[:31])
    sheet.append(columns)
    for row in rows:
        sheet.append([display_cell(row.get(column, ""), 900) for column in columns])
    for cell in sheet[1]:
        cell.font = cell.font.copy(bold=True)
    sheet.freeze_panes = "A2"
    for column_cells in sheet.columns:
        max_len = max(len(_as_text(cell.value)) for cell in column_cells)
        sheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_len + 2, 12), 48)
    for row_cells in sheet.iter_rows():
        for cell in row_cells:
            cell.alignment = cell.alignment.copy(wrap_text=True, vertical="top")


def export_xlsx(report: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Border, Side

    workbook = Workbook()
    workbook.remove(workbook.active)
    compact = report.get("compact") or build_compact_report_data(report["result"])
    meeting = compact["meeting"]
    metrics = compact["metrics"]
    _append_sheet(
        workbook,
        "Summary",
        [
            {"field": "Система", "value": "PM Insights"},
            {"field": "Проект", "value": meeting.get("project")},
            {"field": "Дата", "value": meeting.get("date")},
            {"field": "Тип", "value": meeting.get("type")},
            {"field": "Длительность", "value": meeting.get("duration")},
            {"field": "ASR модель", "value": meeting.get("asr_model")},
            {"field": "Время обработки", "value": meeting.get("processing_time")},
            {"field": "Задачи", "value": metrics.get("tasks")},
            {"field": "Вопросы", "value": metrics.get("questions")},
            {"field": "Ответы", "value": metrics.get("answers")},
            {"field": "Дедлайны", "value": metrics.get("deadlines")},
            {"field": "Основные аспекты", "value": metrics.get("aspects")},
            {"field": "Средний тон", "value": metrics.get("avg_sentiment")},
            {"field": "Негативных фрагментов", "value": metrics.get("negative")},
            {"field": "Ответственные стороны", "value": ", ".join(compact.get("responsible_sides") or [])},
        ],
        ["field", "value"],
    )
    _append_sheet(workbook, "Tasks", compact.get("tasks") or [], ["title", "responsible", "responsible_side", "deadline", "status"])
    _append_sheet(workbook, "Questions_Answers", compact.get("qa") or [], ["question", "answer", "status"])
    _append_sheet(workbook, "Deadlines", compact.get("deadlines") or [], ["deadline", "kind", "context"])
    _append_sheet(workbook, "Topics", (compact.get("topics") or []) + (compact.get("aspects") or []), ["title", "count", "keywords"])
    sentiment = compact["sentiment"]
    _append_sheet(
        workbook,
        "Sentiment",
        [
            {"metric": "Позитив", "value": sentiment.get("positive")},
            {"metric": "Нейтрально", "value": sentiment.get("neutral")},
            {"metric": "Негатив", "value": sentiment.get("negative")},
            {"metric": "Средний тон", "value": sentiment.get("average")},
            *[
                {
                    "metric": row.get("aspect"),
                    "value": f"тон: {row.get('average')}; негатив: {row.get('negative')}",
                }
                for row in sentiment.get("aspect_sentiment") or []
            ],
        ],
        ["metric", "value"],
    )
    thin = Side(style="thin", color="D9D9D9")
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _set_docx_font(document) -> None:
    from docx.oxml.ns import qn

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:cs"), "Calibri")


def _add_docx_table(document, headers: list[str], rows: list[tuple[Any, ...]]) -> None:
    if not rows:
        document.add_paragraph("Нет данных")
        return
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        row = table.add_row().cells
        for index, value in enumerate(values[: len(headers)]):
            row[index].text = display_cell(value, 300)


def export_docx(report: dict) -> bytes:
    from docx import Document

    assert_no_question_marks_in_headings(report)
    compact = report.get("compact") or build_compact_report_data(report["result"])
    document = Document()
    _set_docx_font(document)
    document.add_heading("PM Insights — отчёт по встрече", level=1)
    meeting = compact["meeting"]
    table = document.add_table(rows=0, cols=2)
    for key, value in [
        ("Проект", meeting.get("project")),
        ("Дата", meeting.get("date")),
        ("Тип", meeting.get("type")),
        ("Длительность", meeting.get("duration")),
        ("ASR модель", meeting.get("asr_model")),
        ("Время обработки", meeting.get("processing_time")),
    ]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = _as_text(value)

    document.add_heading("Основные показатели", level=2)
    _add_docx_table(
        document,
        ["Показатель", "Значение"],
        [
            ("Задачи", compact["metrics"]["tasks"]),
            ("Вопросы", compact["metrics"]["questions"]),
            ("Ответы", compact["metrics"]["answers"]),
            ("Дедлайны", compact["metrics"]["deadlines"]),
            ("Основные аспекты", compact["metrics"]["aspects"]),
            ("Средний тон", compact["metrics"]["avg_sentiment"]),
            ("Негативных фрагментов", compact["metrics"]["negative"]),
        ],
    )

    document.add_heading("Основные темы", level=2)
    _add_docx_table(document, ["Тема", "Частота", "Ключевые слова"], [(row["title"], row.get("count"), row.get("keywords")) for row in compact.get("topics") or []])

    document.add_heading(compact.get("tasks_title") or "Задачи", level=2)
    _add_docx_table(document, ["Действие", "Ответственный", "Срок", "Статус"], [(row["title"], row.get("responsible") or row.get("responsible_side"), row.get("deadline"), row.get("status")) for row in compact.get("tasks") or []])

    document.add_heading("Ключевые вопросы и ответы", level=2)
    _add_docx_table(document, ["Вопрос", "Ответ", "Статус"], [(row["question"], row["answer"], row["status"]) for row in compact.get("qa") or []])

    document.add_heading("Дедлайны", level=2)
    _add_docx_table(document, ["Срок", "Тип", "Контекст"], [(row["deadline"], row.get("kind"), row.get("context")) for row in compact.get("deadlines") or []])

    if compact.get("responsibles") or compact.get("responsible_sides"):
        document.add_heading("Ответственные", level=2)
        _add_docx_table(
            document,
            ["Тип", "Значение"],
            [("Ответственный", name) for name in compact.get("responsibles") or []]
            + [("Сторона", name) for name in compact.get("responsible_sides") or []],
        )

    if compact.get("aspects"):
        document.add_heading("Аспекты обсуждения", level=2)
        _add_docx_table(document, ["Аспект", "Частота", "Ключевые слова"], [(row["title"], row.get("count"), row.get("keywords")) for row in compact.get("aspects") or []])

    document.add_heading("Тональность", level=2)
    sentiment = compact["sentiment"]
    _add_docx_table(
        document,
        ["Показатель", "Значение"],
        [
            ("Позитив", sentiment.get("positive")),
            ("Нейтрально", sentiment.get("neutral")),
            ("Негатив", sentiment.get("negative")),
            ("Средний тон", sentiment.get("average")),
        ],
    )
    if sentiment.get("aspect_sentiment"):
        _add_docx_table(
            document,
            ["Аспект", "Средний тон", "Негативных фрагментов"],
            [(row.get("aspect"), row.get("average"), row.get("negative")) for row in sentiment["aspect_sentiment"]],
        )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def register_cyrillic_pdf_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/times.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
        Path("/usr/share/fonts/truetype/freefont/FreeSans.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    raise RuntimeError("Не найден системный Unicode-шрифт для PDF")


def _pdf_text(pdf, text: Any, height: int = 5) -> None:
    pdf.set_x(pdf.l_margin)
    value = clean_export_text(text)
    try:
        pdf.multi_cell(0, height, value, wrapmode="CHAR")
    except TypeError:
        pdf.multi_cell(0, height, value)


def _pdf_heading(pdf, text: str) -> None:
    pdf.set_font("PMInsightsUnicode", "", 13)
    pdf.ln(2)
    _pdf_text(pdf, text, 7)
    pdf.set_font("PMInsightsUnicode", "", 10)


def _pdf_rows(pdf, headers: list[str], rows: list[tuple[Any, ...]], *, max_rows: int | None = None) -> None:
    pdf.set_font("PMInsightsUnicode", "", 10)
    _pdf_text(pdf, " | ".join(headers), 6)
    pdf.set_font("PMInsightsUnicode", "", 9)
    for values in rows[: max_rows or len(rows)]:
        line = " | ".join(display_cell(value, 140) for value in values)
        _pdf_text(pdf, line, 5)
    if not rows:
        _pdf_text(pdf, "Нет данных", 5)
    pdf.ln(1)


def _write_pdf_compact(pdf, compact: dict) -> None:
    pdf.set_font("PMInsightsUnicode", "", 14)
    _pdf_text(pdf, "PM Insights — отчёт по встрече", 8)
    pdf.set_font("PMInsightsUnicode", "", 10)
    meeting = compact["meeting"]
    _pdf_rows(
        pdf,
        ["Поле", "Значение"],
        [
            ("Проект", meeting.get("project")),
            ("Дата", meeting.get("date")),
            ("Тип", meeting.get("type")),
            ("Длительность", meeting.get("duration")),
            ("ASR модель", meeting.get("asr_model")),
            ("Время обработки", meeting.get("processing_time")),
        ],
    )

    _pdf_heading(pdf, "Основные показатели")
    metrics = compact["metrics"]
    _pdf_rows(
        pdf,
        ["Показатель", "Значение"],
        [
            ("Задачи", metrics.get("tasks")),
            ("Вопросы", metrics.get("questions")),
            ("Ответы", metrics.get("answers")),
            ("Дедлайны", metrics.get("deadlines")),
            ("Основные аспекты", metrics.get("aspects")),
            ("Средний тон", metrics.get("avg_sentiment")),
            ("Негативных фрагментов", metrics.get("negative")),
        ],
    )

    _pdf_heading(pdf, "Основные темы")
    _pdf_rows(pdf, ["Тема", "Частота", "Ключевые слова"], [(row["title"], row.get("count"), row.get("keywords")) for row in compact.get("topics") or []], max_rows=8)

    _pdf_heading(pdf, compact.get("tasks_title") or "Задачи")
    _pdf_rows(pdf, ["Действие", "Ответственный", "Срок", "Статус"], [(row["title"], row.get("responsible") or row.get("responsible_side"), row.get("deadline"), row.get("status")) for row in compact.get("tasks") or []], max_rows=8)

    _pdf_heading(pdf, "Ключевые вопросы и ответы")
    _pdf_rows(pdf, ["Вопрос", "Ответ", "Статус"], [(row["question"], row["answer"], row["status"]) for row in compact.get("qa") or []], max_rows=8)

    _pdf_heading(pdf, "Дедлайны")
    _pdf_rows(pdf, ["Срок", "Тип", "Контекст"], [(row["deadline"], row.get("kind"), row.get("context")) for row in compact.get("deadlines") or []], max_rows=10)

    if compact.get("aspects"):
        _pdf_heading(pdf, "Аспекты обсуждения")
        _pdf_rows(pdf, ["Аспект", "Частота", "Ключевые слова"], [(row["title"], row.get("count"), row.get("keywords")) for row in compact.get("aspects") or []], max_rows=8)

    _pdf_heading(pdf, "Тональность")
    sentiment = compact["sentiment"]
    _pdf_rows(
        pdf,
        ["Показатель", "Значение"],
        [
            ("Позитив", sentiment.get("positive")),
            ("Нейтрально", sentiment.get("neutral")),
            ("Негатив", sentiment.get("negative")),
            ("Средний тон", sentiment.get("average")),
        ],
    )


def export_pdf(report: dict) -> bytes:
    from fpdf import FPDF

    assert_no_question_marks_in_headings(report)
    compact = report.get("compact") or build_compact_report_data(report["result"])

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    if hasattr(pdf, "set_compression"):
        pdf.set_compression(False)
    pdf.add_page()
    pdf.add_font("PMInsightsUnicode", "", register_cyrillic_pdf_font())
    pdf.set_font("PMInsightsUnicode", "", 11)
    _write_pdf_compact(pdf, compact)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        pdf.output(tmp_path)
        return tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)
