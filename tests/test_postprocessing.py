from io import BytesIO

from pm_insights.nlp.postprocessing import normalize_analysis_result


def frag(index, text):
    return {"fragment_index": index, "text": text, "start": float(index), "end": float(index + 1)}


def test_clean_task_normalizes_need_to_do_phrase():
    result = {
        "transcript": [frag(1, "надо в этом плане именно развертку сделать.")],
        "tasks": [{"text": "надо в этом плане именно развертку сделать.", "source_fragment": 1, "confidence": 0.8}],
        "questions_answers": [],
        "topics": [{"topic_name": "гидродинамика"}],
        "aspects": [{"aspects": ["гидродинамика"]}],
    }
    normalized = normalize_analysis_result(result)
    assert normalized["clean_tasks"][0]["title"].startswith("Сделать развертку")
    assert normalized["clean_tasks"][0]["review_required"] is True
    assert normalized["tasks"]


def test_clean_task_keeps_parameter_split_action_item():
    text = "Тогда из ближайших задач получается попробовать описанный вариант разделения на S, N, A, L параметры"
    normalized = normalize_analysis_result(
        {
            "transcript": [frag(1, text)],
            "tasks": [{"text": text, "source_fragment": 1, "confidence": 0.7}],
            "questions_answers": [],
            "topics": [{"topic_name": "параметры пласта"}],
            "aspects": [{"aspects": ["параметры пласта"]}],
        }
    )
    titles = [item["title"] for item in normalized["clean_tasks"]]
    assert "Проверить разделение исследования по группам параметров S, N, A, L" in titles


def test_clean_qa_finds_answer_in_following_window():
    result = {
        "transcript": [
            frag(1, "Илья, можно вязкость тоже сделать?"),
            frag(2, "Да, для нефтяного кейса это получается значение отдельное."),
            frag(3, "Если рассматриваем такой вариант, оно идет как параметр флюида."),
        ],
        "tasks": [],
        "questions_answers": [
            {
                "question": "Илья, можно вязкость тоже сделать?",
                "answer": None,
                "status": "not_answered",
                "question_fragment": 1,
            }
        ],
        "topics": [{"topic_name": "вязкость"}],
        "aspects": [{"aspects": ["вязкость"]}],
    }
    normalized = normalize_analysis_result(result)
    pair = normalized["clean_questions_answers"][0]
    assert pair["status"] == "answered"
    assert "нефтяного кейса" in pair["answer"]
    assert pair["source_fragments"] == [1, 2, 3]


def test_technical_summary_and_raw_fields_are_preserved():
    normalized = normalize_analysis_result(
        {
            "transcript": [frag(1, "гидродинамика, дебит, скважина, МГРП, параметры")],
            "tasks": [{"text": "надо в этом плане именно развертку сделать.", "source_fragment": 1}],
            "questions_answers": [],
            "topics": [{"topic_name": "гидродинамика"}, {"topic_name": "ВКР"}],
            "aspects": [{"aspects": ["гидродинамика", "ВКР"]}],
        }
    )
    assert normalized["analysis_summary"]["meeting_type"] == "technical_research"
    assert normalized["analysis_summary"]["requires_manual_review"] is True
    assert normalized["tasks"]
    assert normalized["clean_tasks"]


def test_clean_schema_adds_review_and_quality_fields():
    normalized = normalize_analysis_result(
        {
            "transcript": [frag(1, "Нужно скинуть промежуточные результаты к следующей встрече")],
            "semantic_blocks": [frag(1, "Нужно скинуть промежуточные результаты к следующей встрече")],
            "tasks": [{"text": "Нужно скинуть промежуточные результаты к следующей встрече", "source_fragment": 1}],
            "questions_answers": [],
            "deadlines": [
                {
                    "text": "Нужно скинуть промежуточные результаты к следующей встрече",
                    "deadlines": ["к следующей встрече"],
                    "kind": "task_deadline",
                    "source_fragment": 1,
                }
            ],
            "responsibles": [],
            "decisions": [],
            "topics": [],
            "aspects": [],
        }
    )
    assert "clean_deadlines" in normalized
    assert "clean_responsibles" in normalized
    assert "review_items" in normalized
    assert normalized["quality_metrics"]["clean_tasks_count"] == len(normalized["clean_tasks"])


def test_commercial_oil_gas_postprocessing_compacts_and_filters_tasks():
    result = {
        "transcript": [
            frag(1, "На июль подтверждаем 180 тысяч тонн. Премия по первой партии 1,4 доллара за баррель."),
            frag(2, "Мы можем прислать последние лабораторные данные отдельно."),
            frag(3, "Тогда по платежам нужно прописать точные даты."),
            frag(4, "Документы можем отправить сегодня-завтра."),
            frag(5, "А период ценообразования какой вы хотите?"),
            frag(6, "Мы предпочитаем среднее за 5 котировочных дней вокруг коносамента. Это приемлемо."),
        ],
        "tasks": [
            {"text": "Мы можем прислать последние лабораторные данные отдельно.", "source_fragment": 2},
            {"text": "Тогда по платежам нужно прописать точные даты.", "source_fragment": 3},
            {"text": "Документы можем отправить сегодня-завтра.", "source_fragment": 4},
        ],
        "questions_answers": [
            {
                "question": "А период ценообразования какой вы хотите? Следующая фраза уже не вопрос.",
                "answer": "Мы предпочитаем среднее за 5 котировочных дней вокруг коносамента. Это приемлемо. Тогда хедж тоже проще поставить.",
                "status": "answered",
                "question_fragment": 5,
                "answer_fragment": 6,
            }
        ],
        "deadlines": [
            {
                "text": "Документы можем отправить сегодня-завтра.",
                "deadlines": ["сегодня-завтра"],
                "kind": "mention",
                "source_fragment": 4,
            }
        ],
        "agreements": [
            {
                "text": "Премия по первой партии 1,4 доллара за баррель.",
                "type": "commercial_term",
                "source_fragment": 1,
                "confidence": 0.72,
            }
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    assert normalized["meeting_type"]["label"] == "commercial_oil_gas"
    titles = [item["title"] for item in normalized["clean_tasks"]]
    assert titles == ["Прописать точные даты платежей"]
    assert len(normalized["commitments"]) == 2
    assert normalized["clean_questions_answers"][0]["question_title"].endswith("?")
    assert len(normalized["clean_questions_answers"][0]["answer_summary"]) <= 330
    assert normalized["clean_deadlines"][0]["context"]
    assert normalized["clean_commercial_terms"][0]["category"] == "премия"


def test_universal_report_layer_adds_sections_and_display_config():
    result = {
        "transcript": [
            frag(1, "По контракту нужно прописать точные даты платежей."),
            frag(2, "Покупатель подтверждает оплату, поставщик направляет term sheet."),
        ],
        "tasks": [{"text": "По контракту нужно прописать точные даты платежей.", "source_fragment": 1}],
        "questions_answers": [
            {
                "question": "Когда покупатель подтверждает оплату?",
                "answer": "Покупатель подтверждает оплату после согласования term sheet.",
                "status": "answered",
                "question_fragment": 1,
                "answer_fragment": 2,
            }
        ],
        "deadlines": [
            {
                "text": "Покупатель подтверждает оплату в течение 24 часов.",
                "deadlines": ["в течение 24 часов"],
                "kind": "answer_deadline",
                "source_fragment": 2,
            }
        ],
        "agreements": [
            {
                "text": "Поставщик направляет term sheet, покупатель подтверждает оплату.",
                "type": "commercial_term",
                "source_fragment": 2,
                "confidence": 0.76,
            }
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    assert normalized["meeting_type"]["label"] == "commercial_meeting"
    assert "report_sections" in normalized
    assert "display_config" in normalized
    assert normalized["clean_tasks"][0]["summary"]
    assert normalized["clean_questions_answers"][0]["question_title"].endswith("?")
    assert normalized["clean_deadlines"][0]["deadline"] == "в течение 24 часов"
    section_ids = [section["id"] for section in normalized["report_sections"]]
    assert section_ids.index("commercial_terms") < section_ids.index("tasks")
    assert normalized["display_config"]["transcript_preview_chars"] == 1600


def test_clean_output_filters_intro_weak_qa_and_frequency_deadlines():
    result = {
        "transcript": [
            frag(1, "Доброе утро, коллеги. У нас сегодня основная тема — поставка."),
            frag(2, "На июль фиксируем 180 тысяч тонн, три партии по 60 тысяч тонн."),
            frag(3, "Давайте запишем: инспекцию делим 50 на 50."),
            frag(4, "Насколько небольшой?"),
            frag(5, "Обновления по логистике два раза в неделю, за пять дней до отгрузки — ежедневный статус."),
        ],
        "tasks": [
            {"text": "Доброе утро, коллеги. У нас сегодня основная тема — поставка.", "source_fragment": 1},
        ],
        "questions_answers": [
            {
                "question": "Насколько небольшой?",
                "answer": "Небольшой, это обсуждается отдельно.",
                "status": "answered",
                "question_fragment": 4,
                "answer_fragment": 5,
            }
        ],
        "deadlines": [
            {
                "text": "Обновления по логистике два раза в неделю, за пять дней до отгрузки — ежедневный статус.",
                "deadlines": ["два раза в неделю", "за пять дней до отгрузки"],
                "kind": "mention",
                "source_fragment": 5,
            }
        ],
        "agreements": [
            {
                "text": "Доброе утро, коллеги. У нас сегодня основная тема — поставка.",
                "type": "agreement",
                "source_fragment": 1,
                "confidence": 0.8,
            },
            {
                "text": "На июль фиксируем 180 тысяч тонн, три партии по 60 тысяч тонн.",
                "type": "commercial_term",
                "source_fragment": 2,
                "confidence": 0.82,
            },
            {
                "text": "Давайте запишем: инспекцию делим 50 на 50.",
                "type": "agreement",
                "source_fragment": 3,
                "confidence": 0.82,
            },
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    assert not normalized["clean_tasks"]
    assert all("Доброе утро" not in item.get("title", "") for item in normalized["clean_commercial_terms"])
    assert [item["title"] for item in normalized["clean_commercial_terms"]] == ["График поставки: три партии по 60 тыс. тонн"]
    assert any(item["title"] == "Инспекция: 50/50 по базовой инспекции" for item in normalized["clean_agreements"])
    assert normalized["clean_questions_answers"] == []
    deadline_values = [item["deadline"] for item in normalized["clean_deadlines"]]
    assert "два раза в неделю" not in deadline_values
    assert "за пять дней до отгрузки" in deadline_values
    assert any(item.get("frequency") == "два раза в неделю" for item in normalized["clean_commitments"])


def test_oil_gas_commercial_term_titles_are_specific():
    result = {
        "transcript": [
            frag(1, "Мы бы предлагали три партии по 60 тысяч тонн."),
            frag(2, "Тогда можем подтвердить минимальный августовский объем 150 тысяч тонн, а еще 70 тысяч тонн оставить опционам до 10 июля."),
            frag(3, "По премии корректировки примерно на 1,8 доллара за баррель."),
            frag(4, "По первой партии премия 1,4, вторая и третья плюс 1,2. Если считать по всем трем партиям средняя премия получится около 1,27."),
        ],
        "tasks": [],
        "questions_answers": [],
        "deadlines": [],
        "agreements": [
            {"text": "Мы бы предлагали три партии по 60 тысяч тонн.", "type": "commercial_term", "source_fragment": 1},
            {
                "text": "Тогда можем подтвердить минимальный августовский объем 150 тысяч тонн, а еще 70 тысяч тонн оставить опционам до 10 июля.",
                "type": "commercial_term",
                "source_fragment": 2,
            },
            {"text": "По премии корректировки примерно на 1,8 доллара за баррель.", "type": "commercial_term", "source_fragment": 3},
            {
                "text": "По первой партии премия 1,4, вторая и третья плюс 1,2. Если считать по всем трем партиям средняя премия получится около 1,27.",
                "type": "commercial_term",
                "source_fragment": 4,
            },
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    titles = [item["title"] for item in normalized["clean_commercial_terms"]]
    assert "График поставки: три партии по 60 тыс. тонн" in titles
    assert "Августовский твердый объем: 150 тыс. тонн" in titles
    assert "Августовский опцион: 70 тыс. тонн до 10 июля" in titles
    assert "Исходное предложение по премии: +1,8 доллара/баррель" in titles
    assert "Премия первой партии: +1,4 доллара/баррель" in titles
    assert "Премия второй и третьей партии: +1,2 доллара/баррель" in titles
    assert "Средняя премия по трем партиям: около +1,27 доллара/баррель" in titles
    assert not normalized["clean_agreements"]


def test_oil_gas_source_mapping_uses_specific_august_fragment():
    result = {
        "transcript": [
            frag(1, "Доброе утро, коллеги. У нас сегодня основная тема – объемы на июль и август."),
            frag(2, "На июль можем подтвердить 180 тысяч тонн сырой нефти."),
            frag(3, "На август – ориентировочно 220 тысяч тонн, но есть зависимость от транспортного окна."),
        ],
        "tasks": [],
        "questions_answers": [],
        "deadlines": [],
        "agreements": [
            {
                "text": "Доброе утро, коллеги. У нас сегодня основная тема – объемы на июль и август. На июль можем подтвердить 180 тысяч тонн сырой нефти. На август – ориентировочно 220 тысяч тонн, но есть зависимость от транспортного окна.",
                "type": "commercial_term",
                "source_fragment": 1,
            }
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    august = next(item for item in normalized["clean_commercial_terms"] if item["title"] == "Августовский ориентировочный объем: 220 тыс. тонн")
    assert "На август" in august["source_text"]
    assert "220 тысяч тонн" in august["source_text"]
    assert "На июль" not in august["source_text"]
    assert august["source_fragment"] == 3


def test_oil_gas_premium_terms_use_direct_split_source():
    result = {
        "transcript": [
            frag(
                7,
                "Премию тогда можно разделить. По первой партии премия 1,4, потому что она срочная. По второй и третьей – 1,2. Если считать по всем трем партиям средняя премия получится около 1,27.",
            )
        ],
        "tasks": [],
        "questions_answers": [],
        "deadlines": [],
        "agreements": [
            {
                "text": "Премию тогда можно разделить. По первой партии премия 1,4, потому что она срочная. По второй и третьей – 1,2. Если считать по всем трем партиям средняя премия получится около 1,27.",
                "type": "commercial_term",
                "source_fragment": 7,
            }
        ],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    terms = {item["title"]: item for item in normalized["clean_commercial_terms"]}
    first = terms["Премия первой партии: +1,4 доллара/баррель"]
    second = terms["Премия второй и третьей партии: +1,2 доллара/баррель"]
    for item in (first, second):
        assert "По первой партии премия 1,4" in item["source_text"]
        assert "По второй и третьей" in item["source_text"]
        assert item["source_fragment"] == 7


def test_oil_gas_premium_breakdown_qa_deduplicated_and_summarized():
    result = {
        "transcript": [],
        "tasks": [],
        "questions_answers": [
            {
                "question": "Можете разложить премию 1,8 доллара по фрахту, страховке и риску?",
                "answer": "Около 90 центов относится к фрахту, 30–40 центов — к страховке и портовым расходам, остальное — к риску задержек и доступности логистики.",
                "status": "answered",
                "question_fragment": 1,
                "answer_fragment": 2,
            },
            {
                "question": "Сколько там фрахт, сколько страховка и сколько риск по премии?",
                "answer": "Около 90 центов относится к фрахту, 30–40 центов — к страховке и портовым расходам, остальное — к риску задержек и доступности логистики.",
                "status": "answered",
                "question_fragment": 3,
                "answer_fragment": 4,
            },
        ],
        "deadlines": [],
        "agreements": [],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    premium_pairs = [item for item in normalized["clean_questions_answers"] if "премия 1,8" in item["question_title"]]
    assert len(premium_pairs) == 1
    assert premium_pairs[0]["question_title"] == "Как раскладывается премия 1,8 доллара по фрахту, страховке и риску?"
    assert premium_pairs[0]["answer_summary"] == "Около 90 центов относится к фрахту, 30–40 центов — к страховке и портовым расходам, остальное — к риску задержек и доступности логистики."


def test_oil_gas_deadline_keeps_single_specific_five_working_days_value():
    result = {
        "transcript": [],
        "tasks": [],
        "questions_answers": [],
        "deadlines": [
            {
                "text": "Номинация судна направляется минимум за 5 рабочих дней до отгрузки.",
                "deadlines": ["за 5 рабочих дней до", "минимум за 5 рабочих дней"],
                "kind": "task_deadline",
                "source_fragment": 9,
            }
        ],
        "agreements": [],
        "topics": [],
        "aspects": [],
    }
    normalized = normalize_analysis_result(result)
    values = [item["deadline"] for item in normalized["clean_deadlines"]]
    assert values == ["минимум за 5 рабочих дней до отгрузки"]


def test_oil_gas_topics_aspects_qa_and_deadline_cleanup_regression():
    result = {
        "meeting_type": {"label": "commercial_oil_gas", "confidence": 0.9},
        "transcript": [
            frag(1, "По контракту обсуждали рабочие даты отгрузки, поставки и платежей."),
            frag(2, "И там будет небольшой рост стоимости логистики, потому что придется двигать часть внутреннего графика. Насколько небольшой?"),
            frag(3, "Около 0,3 доллара за баррель относится к переносу логистики."),
            frag(4, "Документы можем отправить сегодня-завтра."),
            frag(5, "Номинация судна направляется минимум за 5 рабочих дней до отгрузки."),
            frag(6, "Покупатель должен подтвердить приемлемость судна в течение 24 часов."),
        ],
        "tasks": [],
        "questions_answers": [
            {
                "question": "И там будет небольшой рост стоимости логистики, потому что придется двигать часть внутреннего графика. Насколько небольшой?",
                "answer": "Около 0,3 доллара за баррель относится к переносу логистики.",
                "status": "answered",
                "question_fragment": 2,
                "answer_fragment": 3,
            }
        ],
        "deadlines": [
            {
                "text": "Документы можем отправить сегодня-завтра.",
                "deadlines": ["сегодня-завтра"],
                "kind": "mention",
                "source_fragment": 4,
            },
            {
                "text": "Номинация судна направляется минимум за 5 рабочих дней до отгрузки.",
                "deadlines": ["за 5 рабочих дней до отгрузки"],
                "kind": "task_deadline",
                "source_fragment": 5,
            },
            {
                "text": "Покупатель должен подтвердить приемлемость судна в течение 24 часов.",
                "deadlines": ["в течение 24 часов"],
                "kind": "answer_deadline",
                "source_fragment": 6,
            },
        ],
        "topics": [
            {
                "topic_name": "Даты рабочих контракте",
                "keywords": ["даты", "рабочих", "контракте"],
                "count": 2,
                "fragment_ids": [1],
            },
            {
                "topic_name": "даты рабочих контракте",
                "keywords": ["рабочих", "контракте", "отгрузки", "поставки", "даты"],
                "count": 3,
                "fragment_ids": [1, 5],
            },
            {
                "topic_name": "рабочих контракте отгрузки поставки даты",
                "keywords": ["рабочих", "контракте", "отгрузки", "поставки", "даты"],
                "count": 1,
                "fragment_ids": [5],
            },
            {
                "topic_name": "даты платежей",
                "keywords": ["даты", "платежей"],
                "count": 1,
                "fragment_ids": [1],
            },
        ],
        "aspect_frequencies": {
            "Даты рабочих контракте": 2,
            "даты рабочих контракте": 3,
        },
        "aspects": [
            {"source_fragment": 1, "aspects": ["Даты рабочих контракте", "даты рабочих контракте"]},
        ],
        "agreements": [],
    }

    normalized = normalize_analysis_result(result)
    topic_names = [item["topic_name"] for item in normalized["clean_topics"]]
    aspect_names = [item["title"] for item in normalized["clean_aspects"]]
    lowered_topics = [name.lower() for name in topic_names]
    lowered_aspects = [name.lower() for name in aspect_names]

    assert "даты рабочих контракте" not in lowered_topics
    assert "даты рабочих контракте" not in lowered_aspects
    assert "сроки поставки и оплаты" in lowered_topics
    assert "сроки поставки и оплаты" in lowered_aspects
    assert lowered_aspects.count("сроки поставки и оплаты") == 1
    assert "платежные условия" in lowered_topics
    assert normalized["clean_questions_answers"][0]["question_title"] == "Насколько небольшой?"

    deadlines_by_value = {item["deadline"]: item["kind"] for item in normalized["clean_deadlines"]}
    assert deadlines_by_value["сегодня-завтра"] == "commitment_deadline"
    assert deadlines_by_value["минимум за 5 рабочих дней до отгрузки"] == "contract_logistics_deadline"
    assert deadlines_by_value["в течение 24 часов"] == "operational_deadline"


def test_technical_research_splits_actions_recommendations_and_notes():
    result = {
        "meeting_type": {"label": "technical_research", "confidence": 0.9},
        "transcript": [
            frag(1, "Метрокубический разделить на сутки – это дебит."),
            frag(2, "Это необходимо сделать."),
            frag(3, "Проверить устойчивость модели и точность."),
            frag(4, "Лучше проверить более частные версии модели."),
            frag(5, "Сделать развертку по параметрам/формуле безразмеривания."),
            frag(6, "Я попробую поискать кейсы."),
        ],
        "tasks": [
            {"text": "Метрокубический разделить на сутки – это дебит.", "source_fragment": 1, "confidence": 0.75},
            {"text": "Это необходимо сделать.", "source_fragment": 2, "confidence": 0.75},
            {"text": "Проверить устойчивость модели и точность.", "source_fragment": 3, "confidence": 0.82},
            {"text": "Лучше проверить более частные версии модели.", "source_fragment": 4, "confidence": 0.74},
            {"text": "Сделать развертку по параметрам/формуле безразмеривания.", "source_fragment": 5, "confidence": 0.82},
            {"text": "Я попробую поискать кейсы.", "source_fragment": 6, "confidence": 0.72},
        ],
        "questions_answers": [],
        "deadlines": [],
        "topics": [{"topic_name": "модель"}],
        "aspects": [{"aspects": ["модель"]}],
    }
    normalized = normalize_analysis_result(result)
    action_titles = [item["title"] for item in normalized["clean_tasks"]]

    assert normalized["clean_research_actions"] == []
    assert normalized["clean_recommendations"] == []
    assert normalized["clean_research_notes"] == []
    assert "Проверить устойчивость модели и точность" in action_titles
    assert "Сделать развертку по параметрам/формуле безразмеривания" in action_titles
    assert not any("Метрокубический" in title for title in action_titles)
    assert not any(title == "Это необходимо сделать" for title in action_titles)
    assert not any("попробую поискать" in title.lower() for title in action_titles)
    assert any(item["reason"] == "too_generic_for_research_action" for item in normalized["review_items"])
    assert any(item["reason"] == "speaker_intent_without_clear_deliverable" for item in normalized["review_items"])


def test_vkr_clean_output_filters_context_topics_and_normalizes_meeting_deadline():
    result = {
        "meeting_type": {"label": "mixed", "confidence": 0.7},
        "transcript": [
            frag(1, "На консультации по ВКР обсуждали черновик и следующую встречу."),
            frag(2, "Это как гипотеза именно того, чтобы проверить как раз-таки этот момент."),
            frag(3, "Я попробую тогда у себя поискать кейсы именно, которые возможно предоставить."),
            frag(4, "Можно будет к этому отталкиваться, но это далеко не бесполезно."),
            frag(5, "Сделать развертку по параметрам/формуле безразмеривания."),
            frag(6, "Разделить исследование по группам S, N, A, L."),
            frag(7, "Давайте следующую встречу поставим на четверг в 7.30. На следующей неделе обсудим правки."),
        ],
        "tasks": [
            {"text": "Это как гипотеза именно того, чтобы проверить как раз-таки этот момент.", "source_fragment": 2, "confidence": 0.75},
            {"text": "Я попробую тогда у себя поискать кейсы именно, которые возможно предоставить.", "source_fragment": 3, "confidence": 0.75},
            {"text": "Можно будет к этому отталкиваться, но это далеко не бесполезно.", "source_fragment": 4, "confidence": 0.75},
            {"text": "Сделать развертку по параметрам/формуле безразмеривания.", "source_fragment": 5, "confidence": 0.82},
            {"text": "Разделить исследование по группам S, N, A, L.", "source_fragment": 6, "confidence": 0.82},
            {"text": "Давайте следующую встречу поставим на четверг в 7.30.", "source_fragment": 7, "confidence": 0.78},
        ],
        "questions_answers": [],
        "deadlines": [
            {
                "text": "Давайте следующую встречу поставим на четверг в 7.30. На следующей неделе обсудим правки.",
                "deadlines": ["четверг", "7.30", "7:30", "следующая неделя", "следующей неделе"],
                "kind": "meeting_time",
                "source_fragment": 7,
            }
        ],
        "topics": [
            {"topic_name": "илья что-нибудь", "keywords": ["илья", "что-нибудь"], "count": 1},
            {"topic_name": "дебет время правило", "keywords": ["дебет", "время", "правило"], "count": 1},
            {"topic_name": "четверг отталкиваться сказали", "keywords": ["четверг", "отталкиваться", "сказали"], "count": 1},
        ],
        "aspects": [],
    }

    normalized = normalize_analysis_result(result)
    task_titles = [item["title"] for item in normalized["clean_tasks"]]
    action_titles = [item["title"] for item in normalized["clean_tasks"]]
    topic_titles = [item["topic_name"].lower() for item in normalized["clean_topics"]]
    deadline_values = [item["deadline"] for item in normalized["clean_deadlines"]]

    assert normalized["meeting_type"]["label"] in {"education_consultation", "technical_research"}
    assert not any("того, чтобы проверить" in title.lower() for title in task_titles)
    assert not any("поискать кейсы" in title.lower() for title in task_titles)
    assert not any("можно будет" in title.lower() for title in task_titles)
    assert "Сделать развертку по параметрам/формуле безразмеривания" in action_titles
    assert "Разделить исследование по группам параметров S, N, A, L" in action_titles
    assert "илья что-нибудь" not in topic_titles
    assert "дебет время правило" not in topic_titles
    assert "четверг отталкиваться сказали" not in topic_titles
    assert "дебит и динамика" in topic_titles
    assert "организация следующей встречи" in topic_titles
    assert "четверг 19:30" in deadline_values
    assert "7.30" not in deadline_values
    assert "7:30" not in deadline_values
    assert deadline_values.count("на следующей неделе") == 1


def test_final_display_quality_gate_filters_bad_topics_and_domain_leaks():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.8},
            "transcript": [
                frag(1, "В ВКР обсуждаем диаметр графа, критерии оптимизации топологии и Raspberry config."),
            ],
            "tasks": [],
            "questions_answers": [],
            "deadlines": [],
            "topics": [
                {"topic_name": "Требования этой эту", "keywords": ["требования", "эту"], "count": 4},
                {"topic_name": "Причем понял честно", "keywords": ["причем", "понял", "честно"], "count": 3},
                {"topic_name": "параметры пласта", "keywords": ["параметры", "пласта"], "count": 2},
                {"topic_name": "Диаметр интересно случайный", "keywords": ["диаметр", "граф", "топология"], "count": 2},
                {"topic_name": "пользоваться config raspberry", "keywords": ["config", "raspberry", "cpu"], "count": 1},
            ],
            "aspects": [
                {"aspects": ["Даты рабочих контракте", "даты рабочих контракте", "параметры пласта"]},
            ],
        }
    )
    topic_names = {item["topic_name"].lower() for item in normalized["clean_topics"]}
    aspect_names = {item["title"].lower() for item in normalized["clean_aspects"]}

    assert "требования этой эту" not in topic_names
    assert "причем понял честно" not in topic_names
    assert "параметры пласта" not in topic_names
    assert "параметры пласта" not in aspect_names
    assert "диаметр и топология графа" in topic_names
    assert "raspberry pi и настройка окружения" in topic_names
    assert "даты рабочих контракте" not in aspect_names


def test_final_clean_filters_qa_deadlines_and_task_noise():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.8},
            "transcript": [
                frag(1, "Что-то сделать, поэтому это вам, видимо, надо будет?"),
                frag(2, "Какие метрики использовать для оценки топологии?"),
                frag(3, "Можно связать диаметр графа, задержку и количество связей."),
                frag(4, "Нужно описать запуск сервиса через systemd timer каждые 15 минут."),
            ],
            "tasks": [
                {"text": "Согласовать следующую встречу на четверг 19:30", "source_fragment": 1, "confidence": 0.8},
                {"text": "Сопоставить кейсы/реализации расчетов", "source_fragment": 2, "confidence": 0.8},
                {"text": "Которую мы писали скинь пожалуйста", "source_fragment": 3, "confidence": 0.8},
                {"text": "Нужно описать запуск сервиса через systemd timer каждые 15 минут.", "source_fragment": 4, "confidence": 0.84},
            ],
            "questions_answers": [
                {"question": "Что-то сделать, поэтому это вам, видимо, надо будет?", "answer": "Да.", "status": "answered", "question_fragment": 1},
                {"question": "Какие метрики использовать для оценки топологии?", "answer": "Диаметр графа, задержку и количество связей.", "status": "answered", "question_fragment": 2},
            ],
            "deadlines": [
                {"text": "Просто срок.", "deadlines": ["срок"], "source_fragment": 1},
                {"text": "В четверг что-то обсудим.", "deadlines": ["четверг"], "source_fragment": 2},
                {"text": "Документ подготовить до 21 числа.", "deadlines": ["до 21 числа"], "kind": "task_deadline", "source_fragment": 4},
            ],
            "topics": [{"topic_name": "systemd timer", "keywords": ["systemd", "timer"]}],
            "aspects": [],
        }
    )

    task_titles = [item["title"] for item in normalized["clean_tasks"]]
    question_titles = [item["question_title"] for item in normalized["clean_questions_answers"]]
    deadline_values = [item["deadline"] for item in normalized["clean_deadlines"]]

    assert "Согласовать следующую встречу на четверг 19:30" not in task_titles
    assert "Сопоставить кейсы/реализации расчетов" not in task_titles
    assert not any("Которую мы писали" in title for title in task_titles)
    assert any("systemd timer" in title.lower() for title in task_titles)
    assert question_titles == ["Какие метрики использовать для оценки топологии?"]
    assert "срок" not in deadline_values
    assert "четверг" not in deadline_values
    assert "до 21 числа" in deadline_values


def test_technical_report_sections_prioritize_research_protocol_blocks():
    result = {
        "meeting_type": {"label": "education_consultation", "confidence": 0.9},
        "transcript": [frag(1, "Проверить устойчивость модели и точность.")],
        "tasks": [{"text": "Проверить устойчивость модели и точность.", "source_fragment": 1, "confidence": 0.82}],
        "questions_answers": [],
        "deadlines": [],
        "topics": [{"topic_name": "ВКР"}],
        "aspects": [{"aspects": ["ВКР"]}],
    }
    normalized = normalize_analysis_result(result)
    section_ids = [section["id"] for section in normalized["report_sections"]]
    assert "tasks" in section_ids
    assert "research_actions" not in section_ids
    assert "recommendations" not in section_ids
    assert "research_notes" not in section_ids
    assert section_ids.index("tasks") < section_ids.index("qa")


def test_normalized_result_adds_universal_clean_layers():
    from pm_insights.nlp.postprocessing import normalize_analysis_result

    result = normalize_analysis_result(
        {
            "meeting_type": {"label": "project_meeting", "confidence": 0.8},
            "topics": [
                {
                    "topic_name": "релиз",
                    "keywords": ["релиз", "срок"],
                    "fragment_ids": [1],
                    "count": 1,
                }
            ],
            "aspects": [{"source_fragment": 1, "aspects": ["релиз"]}],
            "sentiment": [{"source_fragment": 1, "text": "Релиз идет нормально.", "sentiment": "positive", "score": 0.7}],
            "transcript": [{"fragment_index": 1, "text": "Релиз идет нормально."}],
        }
    )
    assert result["clean_topics"][0]["title"] == "Релиз"
    assert result["clean_aspects"][0]["title"] == "Релиз"
    assert result["clean_sentiment"][0]["sentiment"] == "positive"
    assert result["sentiment_summary"]["positive_count"] == 1
    assert "aspect_sentiment" in result
    assert result["clean_notes"] == result["clean_research_notes"]


def test_clean_tasks_use_metadata_participants_as_responsible_gazetteer():
    result = normalize_analysis_result(
        {
            "metadata": {
                "meeting_info": {
                    "participants": [
                        {"name": "Анна", "role": "аналитик"},
                        {"name": "Иван", "role": "разработчик"},
                        {"name": "Мария", "role": "тестировщик"},
                    ]
                }
            },
            "transcript": [
                frag(1, "Анна, подготовь отчёт до пятницы"),
                frag(2, "Иван исправь ошибку авторизации"),
                frag(3, "Мария, проверь релизную сборку"),
            ],
            "tasks": [
                {"text": "Анна, подготовь отчёт до пятницы", "source_fragment": 1},
                {"text": "Иван исправь ошибку авторизации", "source_fragment": 2},
                {"text": "Мария, проверь релизную сборку", "source_fragment": 3},
            ],
            "questions_answers": [],
            "topics": [{"topic_name": "релиз"}],
            "aspects": [{"aspects": ["релиз"]}],
        }
    )
    responsibles_by_fragment = {
        item["source_fragment"]: item["responsible"]
        for item in result["clean_tasks"]
    }
    assert responsibles_by_fragment[1] == "Анна"
    assert responsibles_by_fragment[2] == "Иван"
    assert responsibles_by_fragment[3] == "Мария"


def test_raspberry_yolo_topics_do_not_leak_oil_gas_or_filler_titles():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(
                    1,
                    "Для ВКР обсуждали Raspberry Pi, YOLO benchmark, камеры, QEMU, "
                    "датасет, OpenCV и задачи для второкурсников.",
                )
            ],
            "tasks": [],
            "questions_answers": [],
            "deadlines": [],
            "topics": [
                {"topic_name": "Скин-фактор и трещины", "keywords": ["скин-фактор", "трещины"], "count": 3},
                {"topic_name": "Параметры пласта", "keywords": ["параметры", "пласта"], "count": 3},
                {"topic_name": "пользоваться config raspberry", "keywords": ["config", "raspberry", "cpu"], "count": 3},
                {"topic_name": "пакете пакете", "keywords": ["пакете", "пакете"], "count": 2},
                {"topic_name": "второкурсниками много", "keywords": ["второкурсниками", "много"], "count": 2},
                {"topic_name": "камеры fps", "keywords": ["камеры", "fps"], "count": 2},
                {"topic_name": "qemu виртуализация", "keywords": ["qemu", "виртуализация"], "count": 2},
                {"topic_name": "влезет смысле", "keywords": ["влезет", "смысле"], "count": 2},
                {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 4},
            ],
            "aspects": [
                {
                    "text": "Raspberry Pi, YOLO benchmark, камеры, QEMU, датасет, OpenCV.",
                    "aspects": [
                            "Скин-фактор и трещины",
                            "Параметры пласта",
                            "пользоваться config raspberry",
                            "пакете пакете",
                            "второкурсниками много",
                            "камеры fps",
                            "qemu виртуализация",
                            "влезет смысле",
                        ],
                    }
            ],
        }
    )

    topic_names = {item["topic_name"].lower() for item in normalized["clean_topics"]}
    aspect_names = {item["title"].lower() for item in normalized["clean_aspects"]}
    combined = topic_names | aspect_names

    assert "скин-фактор и трещины" not in combined
    assert "параметры пласта" not in combined
    assert "пакете пакете" not in combined
    assert "второкурсниками много" not in combined
    assert "влезет смысле" not in combined
    assert "raspberry pi и настройка окружения" in combined
    assert "yolo benchmark" in combined
    assert "камеры и fps" in combined
    assert "задачи для второкурсников" in combined
    assert "qemu / виртуализация raspberry pi" in combined


def test_raspberry_yolo_semantic_layer_extracts_actions_deadlines_and_filters_qa():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "В следующий раз приходим с термопастой и проверяем температуру Raspberry Pi."),
                frag(2, "Нужно запустить YOLO benchmark на Raspberry Pi и посмотреть частоту во время benchmark."),
                frag(3, "Датасеты надо привести к общему формату, а камеры проверить в коробках и пакетах."),
                frag(4, "Найти инструкции по AI Kit Hailo и поручить второкурсникам изучить QEMU симуляцию Raspberry Pi."),
                frag(5, "Формальную постановку задачи отправить завтра."),
            ],
            "tasks": [],
            "questions_answers": [
                {
                    "question": "Что можно дать какие-нибудь простенькие задачи второкурсникам, чтобы они не простаивали?",
                    "answer": "Можно QEMU и симуляцию Raspberry Pi.",
                    "status": "answered",
                    "question_fragment": 4,
                },
                {
                    "question": "Можно ли использовать QEMU для симуляции Raspberry Pi?",
                    "answer": "Да, это подойдет для первичного знакомства.",
                    "status": "answered",
                    "question_fragment": 4,
                },
            ],
            "deadlines": [],
            "topics": [
                {"topic_name": "премия и дифференциал", "keywords": ["премия", "дифференциал"], "count": 3},
                {"topic_name": "Скин-фактор и трещины", "keywords": ["скин-фактор"], "count": 2},
                {"topic_name": "пакете пакете", "keywords": ["пакете"], "count": 2},
                {"topic_name": "второкурсниками много", "keywords": ["второкурсниками"], "count": 2},
                {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 5},
            ],
            "aspects": [{"aspects": ["премия и дифференциал", "пакете пакете", "YOLO benchmark"]}],
        }
    )

    topic_names = {item["topic_name"].lower() for item in normalized["clean_topics"]}
    aspect_names = {item["title"].lower() for item in normalized["clean_aspects"]}
    task_titles = {item["title"] for item in normalized["clean_tasks"]}
    deadline_values = {item["deadline"] for item in normalized["clean_deadlines"]}
    question_titles = [item["question_title"] for item in normalized["clean_questions_answers"]]

    assert "премия и дифференциал" not in topic_names | aspect_names
    assert "скин-фактор и трещины" not in topic_names | aspect_names
    assert "пакете пакете" not in topic_names | aspect_names
    assert "второкурсниками много" not in topic_names | aspect_names
    assert "Запустить YOLO benchmark на Raspberry Pi" in task_titles
    assert "Привести датасеты к общему формату" in task_titles
    assert "Поручить второкурсникам изучить QEMU / симуляцию Raspberry Pi" in task_titles
    assert "завтра" in deadline_values
    assert "в следующий раз" in deadline_values
    assert question_titles == [
        "Что можно поручить второкурсникам?",
        "Можно ли использовать QEMU для симуляции Raspberry Pi?",
    ]


def test_raspberry_yolo_export_payload_uses_only_clean_semantic_topics():
    from docx import Document
    from openpyxl import load_workbook

    from pm_insights.export.report_builder import build_report_payload, export_docx, export_pdf, export_xlsx

    normalized = normalize_analysis_result(
        {
            "metadata": {"meeting_info": {"meeting_title": "Raspberry YOLO ВКР", "project_name": "ВКР"}},
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "Для ВКР обсуждали Raspberry Pi и настройку окружения, YOLO benchmark, камеры и FPS."),
                frag(2, "Нужно проверить камеры в коробках и пакетах, QEMU виртуализацию Raspberry Pi и задачи для второкурсников."),
            ],
            "tasks": [],
            "questions_answers": [],
            "deadlines": [],
            "topics": [
                {"topic_name": "Скин-фактор и трещины", "keywords": ["скин-фактор"], "count": 3},
                {"topic_name": "Параметры пласта", "keywords": ["пласт"], "count": 3},
                {"topic_name": "пакете пакете", "keywords": ["пакете"], "count": 2},
                {"topic_name": "второкурсниками много", "keywords": ["второкурсниками"], "count": 2},
                {"topic_name": "пользоваться config raspberry", "keywords": ["config", "raspberry"], "count": 2},
                {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 4},
                {"topic_name": "камеры fps", "keywords": ["камеры", "fps"], "count": 3},
                {"topic_name": "qemu виртуализация", "keywords": ["qemu", "виртуализация"], "count": 3},
            ],
            "aspects": [
                {
                    "aspects": [
                        "Скин-фактор и трещины",
                        "Параметры пласта",
                        "пакете пакете",
                        "второкурсниками много",
                        "пользоваться config raspberry",
                        "YOLO benchmark",
                        "камеры fps",
                        "qemu виртуализация",
                    ]
                }
            ],
        }
    )
    report = build_report_payload(normalized)
    bad = ("Скин-фактор и трещины", "Параметры пласта", "пакете пакете", "второкурсниками много")
    good = (
        "Raspberry Pi и настройка окружения",
        "YOLO benchmark",
        "камеры и FPS",
        "задачи для второкурсников",
        "QEMU / виртуализация Raspberry Pi",
    )

    report_text = str(report)
    for value in bad:
        assert value.lower() not in report_text.lower()
    for value in good:
        assert value.lower() in report_text.lower()

    document = Document(BytesIO(export_docx(report)))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            docx_text += "\n" + "\t".join(cell.text for cell in row.cells)
    for value in bad:
        assert value.lower() not in docx_text.lower()
    for value in good:
        assert value.lower() in docx_text.lower()

    workbook = load_workbook(BytesIO(export_xlsx(report)))
    xlsx_text = ""
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            xlsx_text += "\n" + "\t".join("" if value is None else str(value) for value in row)
    for value in bad:
        assert value.lower() not in xlsx_text.lower()
    for value in good:
        assert value.lower() in xlsx_text.lower()

    pdf_bytes = export_pdf(report)
    assert pdf_bytes.startswith(b"%PDF")
    assert len(pdf_bytes) > 1000


def test_c4_architecture_semantic_layer_prefers_architecture_topics_and_actions():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.86},
            "transcript": [
                frag(1, "По C4-диаграмме нужно описать компоненты и взаимодействия через каталог LDAP."),
                frag(2, "Диаграмму лучше отправить отдельным PNG PDF SVG файлом и вставить в ВКР и презентацию."),
                frag(3, "Подпись systemd timer надо уточнить как системный вызов, а термин инфраструктурный сервис проверить."),
                frag(4, "Во вторник покажем архитектурную диаграмму отдельной страницей landscape, если она плохо читается."),
            ],
            "tasks": [],
            "questions_answers": [
                {
                    "question": "Где в котором есть таймер, что администратор целой портала?",
                    "answer": "Это сырой фрагмент.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "Как подписать systemd timer на диаграмме?",
                    "answer": "Как системный вызов или планировщик.",
                    "status": "answered",
                    "question_fragment": 3,
                },
            ],
            "deadlines": [],
            "topics": [
                {"topic_name": "данные и выборка", "keywords": ["данные", "каталог"], "count": 5},
                {"topic_name": "Интерпретация данных", "keywords": ["c4", "компоненты"], "count": 5},
                {"topic_name": "точность модели", "keywords": ["модель"], "count": 4},
                {"topic_name": "systemd timer", "keywords": ["systemd", "timer"], "count": 3},
            ],
            "aspects": [{"aspects": ["данные и выборка", "Интерпретация данных", "systemd timer"]}],
        }
    )

    topic_names = [item["topic_name"] for item in normalized["clean_topics"]]
    topic_names_lower = {name.lower() for name in topic_names}
    task_titles = {item["title"] for item in normalized["clean_tasks"]}
    deadline_values = {item["deadline"] for item in normalized["clean_deadlines"]}

    assert "данные и выборка" not in topic_names_lower
    assert "интерпретация данных" not in topic_names_lower
    assert "точность модели" not in topic_names_lower
    assert "компоненты и взаимодействия" in topic_names_lower or "внешние системы и пользователи" in topic_names_lower
    assert "systemd timer" in topic_names_lower
    assert "Отправить диаграмму отдельным PNG/PDF/SVG-файлом для проверки" in task_titles
    assert "Описать компоненты и взаимодействия по C4-диаграмме" in task_titles
    assert "во вторник" in deadline_values


def test_graph_vkr_semantic_layer_filters_bad_tasks_questions_and_filler_topics():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "Продолжить главу с постановкой задачи и переформулировать критерий оптимальности топологии графа."),
                frag(2, "Нужно описать ограничение по количеству связей, обосновать выбор топологии графа и описать метрики графа."),
                frag(3, "Описать сбор данных и телеметрию, исключения и ограничения модели, критерии качества модели."),
                frag(4, "Найти похожие ВКР и шаблоны оформления."),
            ],
            "tasks": [{"text": "Описать вариант в ВКР", "source_fragment": 1, "confidence": 0.8}],
            "questions_answers": [
                {
                    "question": "Что как бы написать математическую постановку, что значит вот минимизировать и максимизировать?",
                    "answer": "Сырой ASR-фрагмент.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "Какие метрики использовать для оценки топологии?",
                    "answer": "Диаметр графа, задержку и количество связей.",
                    "status": "answered",
                    "question_fragment": 2,
                },
            ],
            "deadlines": [],
            "topics": [
                {"topic_name": "входные и выходные данные модулей", "keywords": ["граф", "параметры"], "count": 5},
                {"topic_name": "Модель и аппроксимация", "keywords": ["модель", "граф"], "count": 4},
                {"topic_name": "ВКР и документация", "keywords": ["вкр"], "count": 3},
                {"topic_name": "Требования этой эту", "keywords": ["эта", "эту"], "count": 4},
            ],
            "aspects": [{"aspects": ["входные и выходные данные модулей", "Требования этой эту"]}],
        }
    )

    task_titles = {item["title"] for item in normalized["clean_tasks"]}
    question_titles = [item["question_title"] for item in normalized["clean_questions_answers"]]
    topic_names = {item["topic_name"].lower() for item in normalized["clean_topics"]}
    aspect_names = {item["title"].lower() for item in normalized["clean_aspects"]}

    assert "Описать вариант в ВКР" not in task_titles
    assert "Переформулировать критерий оптимальности топологии" in task_titles
    assert "Описать ограничение по количеству связей" in task_titles
    assert question_titles == ["Какие метрики использовать для оценки топологии?"]
    assert "требования этой эту" not in topic_names | aspect_names
    assert "входные данные математической модели" in topic_names | aspect_names
    assert "математическая модель" in topic_names | aspect_names


def test_graph_vkr_clean_qa_normalizes_raw_metric_and_telemetry_questions():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "Обсуждаем математическую модель графа, диаметр, среднюю дистанцию, метрики графа и критерии качества."),
                frag(2, "Также нужен сбор данных и телеметрия для проверки модели топологии."),
            ],
            "tasks": [],
            "questions_answers": [
                {
                    "question": "что она ухудшается но не так резко как просто диаметре и может быть смешанная метрика?",
                    "answer": "Да, можно ввести смешанную метрику для оценки графа.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "плохой тем что он учитывает только наихудший случай, может лучше средняя дистанция?",
                    "answer": "Нужно сравнить диаметр и среднюю дистанцию.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "какие-то инструменты для этого нужны?",
                    "answer": "Да, нужно добавить сбор данных и телеметрию.",
                    "status": "answered",
                    "question_fragment": 2,
                },
                {
                    "question": "что нам интересно от этой модели это какие критерии качества?",
                    "answer": "В критериях качества указать диаметр, задержку и количество связей.",
                    "status": "answered",
                    "question_fragment": 1,
                },
            ],
            "deadlines": [],
            "topics": [{"topic_name": "Модель и аппроксимация", "keywords": ["граф", "метрики"], "count": 4}],
            "aspects": [{"aspects": ["Модель и аппроксимация"]}],
        }
    )

    question_titles = [item["question_title"] for item in normalized["clean_questions_answers"]]
    question_blob = " ".join(question_titles).lower()

    assert "что она ухудшается но не так резко" not in question_blob
    assert "плохой тем что он учитывает" not in question_blob
    assert "что нам интересно от этой модели" not in question_blob
    assert "какие-то инструменты для" not in question_blob
    assert "Есть ли смешанная метрика для оценки графа?" in question_titles
    assert "Какие метрики лучше использовать: диаметр или среднюю дистанцию?" in question_titles
    assert "Нужно ли добавить сбор данных и телеметрию?" in question_titles
    assert "Что указать в критериях качества?" in question_titles


def test_graph_vkr_topic_labels_and_aspect_order_are_precise():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "Постановка задачи: описать математическую модель графа и критерии качества."),
                frag(2, "Аппроксимация метрик графа нужна для сравнения диаметра и средней дистанции."),
                frag(3, "Входные данные математической модели остаются отдельным, но вторичным описанием."),
            ],
            "tasks": [],
            "questions_answers": [],
            "deadlines": [],
            "topics": [
                {"topic_name": "входные и выходные данные модулей", "keywords": ["данные", "модель"], "count": 10},
                {"topic_name": "Модель и аппроксимация", "keywords": ["аппроксимация", "метрики", "граф"], "count": 8},
                {"topic_name": "описание архитектуры в ВКР", "keywords": ["вкр"], "count": 7},
                {"topic_name": "постановка задачи", "keywords": ["постановка", "задачи"], "count": 5},
            ],
            "aspects": [
                {
                    "aspects": [
                        "входные и выходные данные модулей",
                        "Модель и аппроксимация",
                        "описание архитектуры в ВКР",
                        "постановка задачи",
                    ]
                }
            ],
        }
    )

    topic_names = [item["topic_name"].lower() for item in normalized["clean_topics"]]
    aspect_names = [item["title"].lower() for item in normalized["clean_aspects"]]

    assert "описание архитектуры в вкр" not in topic_names + aspect_names
    assert "аппроксимация метрик графа" in topic_names + aspect_names
    assert "входные данные математической модели" in topic_names + aspect_names
    assert topic_names.index("постановка задачи") < topic_names.index("входные данные математической модели")
    assert aspect_names.index("постановка задачи") < aspect_names.index("входные данные математической модели")


def test_historical_dates_are_not_clean_deadlines_but_future_actions_are_kept():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "В общем, мы ходили в понедельник, смотрели там, разбирались с Raspberry Pi."),
                frag(2, "Завтра снова прийти с термопастой и отправить формальную постановку задачи."),
                frag(3, "В пятницу будем работать с данными, а к выходным привести датасеты к общему виду."),
                frag(4, "В следующий раз прийти с камерами, через месяц будут ожидаемые данные."),
            ],
            "tasks": [],
            "questions_answers": [],
            "deadlines": [],
            "topics": [{"topic_name": "Raspberry Pi и настройка окружения", "keywords": ["raspberry"]}],
            "aspects": [{"aspects": ["Raspberry Pi и настройка окружения"]}],
        }
    )

    values = {item["deadline"] for item in normalized["clean_deadlines"]}
    assert "в понедельник" not in values
    assert "завтра" in values
    assert "в пятницу" in values
    assert "к выходным" in values
    assert "в следующий раз" in values
    assert "через месяц" in values


def test_raspberry_topics_keywords_and_qa_are_domain_specific():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "Raspberry Pi, YOLO benchmark, камеры FPS, OpenCV, датасеты, Hailo NPU и формальная постановка задачи."),
                frag(2, "В следующий раз план работы с оборудованием: проверить benchmark и камеры."),
            ],
            "tasks": [],
            "questions_answers": [
                {
                    "question": "или что у вас там есть?",
                    "answer": "Есть Raspberry Pi, камеры и датасеты.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "какие-то учебные вещи на opencv можно дать?",
                    "answer": "Можно поручить второкурсникам простые задачи по OpenCV.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "Можно ли смотреть частоту во время benchmark?",
                    "answer": "Да, частоту и температуру стоит смотреть.",
                    "status": "answered",
                    "question_fragment": 2,
                },
            ],
            "deadlines": [],
            "topics": [
                {"topic_name": "входные данные математической модели", "keywords": ["влезет смысл", "какой-нибудь"], "count": 8},
                {"topic_name": "Организация следующей встречи", "keywords": ["следующий", "раз"], "count": 7},
                {"topic_name": "модель топологии сети", "keywords": ["модель", "топология"], "count": 6},
                {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 5},
                {"topic_name": "камеры fps", "keywords": ["камеры", "fps"], "count": 4},
            ],
            "aspects": [{"aspects": ["входные данные математической модели", "Организация следующей встречи", "модель топологии сети", "YOLO benchmark"]}],
        }
    )

    topic_names = [item["topic_name"].lower() for item in normalized["clean_topics"]]
    aspect_names = [item["title"].lower() for item in normalized["clean_aspects"]]
    keywords = " ".join(" ".join(item.get("keywords") or []) for item in normalized["clean_topics"] + normalized["clean_aspects"]).lower()
    questions = [item["question_title"] for item in normalized["clean_questions_answers"]]

    assert "входные данные математической модели" not in topic_names + aspect_names
    assert "организация следующей встречи" not in topic_names + aspect_names
    assert "модель топологии сети" not in topic_names + aspect_names
    assert "yolo benchmark" in topic_names + aspect_names
    assert "камеры и fps" in topic_names + aspect_names
    assert "входные данные ml pipeline" in topic_names + aspect_names or "формальная постановка задачи" in topic_names + aspect_names
    assert "влезет" not in keywords
    assert "смысл" not in keywords
    assert "Что вообще у вас есть?" in questions
    assert "Что можно поручить второкурсникам?" in questions
    assert "Можно ли смотреть частоту во время benchmark?" in questions


def test_c4_topics_qa_and_deadlines_do_not_use_math_placeholders():
    normalized = normalize_analysis_result(
        {
            "meeting_type": {"label": "education_consultation", "confidence": 0.9},
            "transcript": [
                frag(1, "C4 архитектурная диаграмма, systemd timer, LDAP каталог домена, компоненты и взаимодействия."),
                frag(2, "Во вторник нужно прийти с презентацией и отдельным PNG PDF SVG файлом диаграммы."),
                frag(3, "Можно добавить sequence diagram и пояснить каждую стрелку."),
            ],
            "tasks": [],
            "questions_answers": [
                {
                    "question": "что думаете в целом понятно вам как из этого всего дела тут вообще что происходит?",
                    "answer": "Диаграмму надо сделать понятнее.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "какие-то особые дефолтные так получается уровни c4 нужны?",
                    "answer": "Можно не делать лишние уровни C4.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "что там используют какие-то локальные файлы?",
                    "answer": "Локальные артефакты лучше указать отдельно.",
                    "status": "answered",
                    "question_fragment": 1,
                },
                {
                    "question": "Можно ли прислать диаграмму отдельным PNG/PDF/SVG-файлом?",
                    "answer": "Да.",
                    "status": "answered",
                    "question_fragment": 2,
                },
            ],
            "deadlines": [],
            "topics": [
                {"topic_name": "входные данные математической модели", "keywords": ["данные", "модель"], "count": 10},
                {"topic_name": "Организация следующей встречи", "keywords": ["вторник"], "count": 9},
                {"topic_name": "данные и выборка", "keywords": ["данные"], "count": 8},
                {"topic_name": "Интерпретация данных", "keywords": ["данные"], "count": 7},
                {"topic_name": "архитектурная диаграмма C4", "keywords": ["c4", "диаграмма"], "count": 6},
                {"topic_name": "sequence diagram", "keywords": ["sequence"], "count": 5},
            ],
            "aspects": [{"aspects": ["входные данные математической модели", "Организация следующей встречи", "данные и выборка", "архитектурная диаграмма C4"]}],
        }
    )

    names = [item["topic_name"].lower() for item in normalized["clean_topics"]]
    aspect_names = [item["title"].lower() for item in normalized["clean_aspects"]]
    questions = [item["question_title"] for item in normalized["clean_questions_answers"]]
    deadlines = {item["deadline"] for item in normalized["clean_deadlines"]}

    assert "входные данные математической модели" not in names + aspect_names
    assert "организация следующей встречи" not in names + aspect_names
    assert "данные и выборка" not in names + aspect_names
    assert "интерпретация данных" not in names + aspect_names
    assert names[0] == "архитектурная диаграмма c4"
    assert "sequence diagram" in names + aspect_names
    assert "Можно ли прислать диаграмму отдельным PNG/PDF/SVG-файлом?" in questions
    assert "Нужно ли делать отдельные уровни C4?" in questions
    assert "Нужно ли указывать локальные файлы в описании архитектуры?" in questions
    assert "во вторник" in deadlines
