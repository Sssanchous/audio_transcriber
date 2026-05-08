from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


TRANSCRIPTS_DIR = Path("transcripts")
DATASETS_DIR = Path("datasets")
OUTPUT_PATH = DATASETS_DIR / "prelabeled.jsonl"

RAW_JSONL_FILES = [
    Path("dataset_from_protocols.jsonl"),
    DATASETS_DIR / "raw.jsonl",
    DATASETS_DIR / "prelabeled_raw.jsonl",
    DATASETS_DIR / "train.jsonl",
    DATASETS_DIR / "val.jsonl",
]

TRANSCRIPTS_DIR.mkdir(exist_ok=True)
DATASETS_DIR.mkdir(exist_ok=True)

MIN_LEN = 8
MIN_WORDS = 2

FILLERS = {
    "ну", "вот", "как", "бы", "типа", "короче", "значит", "просто",
    "получается", "наверное", "наверно", "вообще", "собственно",
    "в принципе", "там", "тут", "это", "этот", "эта", "эти",
}

QUESTION_WORDS = (
    "кто", "что", "где", "когда", "почему", "зачем", "как",
    "какой", "какая", "какие", "сколько", "чей", "чья", "чьи",
    "можно ли", "нужно ли", "успеем ли", "надо ли", "в чём",
    "в чем", "что за", "а что", "куда", "откуда",
)

IMPERATIVE_WORDS = {
    "сделай", "подготовь", "проверь", "отправь", "создай", "обнови", "исправь",
    "доделай", "согласуй", "закрой", "напиши", "собери", "заверши", "вынеси",
    "выгрузи", "дополни", "перепроверь", "оформи", "протестируй", "настрой",
    "опиши", "сверь", "проведи", "составь", "размести", "изучи", "добавь",
    "переделай", "найди", "уточни", "сократи", "переформулируй",
}

ACTION_INF = {
    "сделать", "подготовить", "проверить", "отправить", "создать", "обновить",
    "исправить", "доделать", "согласовать", "закрыть", "написать", "собрать",
    "завершить", "вынести", "выгрузить", "дополнить", "перепроверить", "оформить",
    "протестировать", "настроить", "описать", "сверить", "разместить", "изучить",
    "начать", "составить", "назначить", "просмотреть", "найти", "добавить",
    "реализовать", "разработать", "отобразить", "сохранить", "сократить",
    "переформулировать", "переделать", "подписать", "задокументировать",
    "провести", "уточнить", "подобрать", "переписать", "исследовать",
}

PROTOCOL_TASK_PATTERNS = [
    r"^\s*назначить\b",
    r"^\s*подготовить\b",
    r"^\s*проверить\b",
    r"^\s*отправить\b",
    r"^\s*создать\b",
    r"^\s*обновить\b",
    r"^\s*исправить\b",
    r"^\s*согласовать\b",
    r"^\s*закрыть\b",
    r"^\s*вынести\b",
    r"^\s*оформить\b",
    r"^\s*изучить\b",
    r"^\s*составить\b",
    r"^\s*просмотреть\b",
    r"^\s*найти\b",
    r"^\s*начать\b",
    r"^\s*подготовка\b",
    r"^\s*провести\b",
    r"^\s*реализовать\b",
    r"^\s*разработать\b",
    r"^\s*доработать\b",
    r"^\s*добавить\b",
    r"^\s*описать\b",
    r"^\s*уточнить\b",
    r"^\s*задокументировать\b",
    r"^\s*протестировать\b",
    r"^\s*студентам\b.+\b(оформить|согласовать|подготовить|составить|сделать|проверить)\b",
    r"^\s*руководител\w*\b.+\b(выполнит|подготовит|проверит|согласует|сделает)\b",
    r"^\s*принять\s+в\s+работу\b",
    r"^\s*вести\s+регулярные\b",
    r"^\s*на\s+следующей\s+неделе\b.+\b(изучить|найти|начать|подготовить|сделать|проверить)\b",
    r"\bнеобходимо\b.+\b(подготовить|оформить|проверить|согласовать|создать|разработать|реализовать|исправить|добавить|описать|уточнить)\b",
    r"\bследует\b.+\b(подготовить|оформить|проверить|согласовать|создать|разработать|реализовать|исправить|добавить|описать|уточнить)\b",
    r"\bнужно\b.+\b(подготовить|оформить|проверить|согласовать|создать|разработать|реализовать|исправить|добавить|описать|уточнить|сделать|посмотреть)\b",
    r"\bнадо\b.+\b(подготовить|оформить|проверить|согласовать|создать|разработать|реализовать|исправить|добавить|описать|уточнить|сделать|посмотреть)\b",
    r"\bпланируется\b.+\b(подготовить|создать|разработать|реализовать|оформить|сделать)\b",
    r"\bпринято\s+решение\b.+\b(подготовить|создать|разработать|оформить|согласовать|добавить|исправить)\b",
    r"\bрекомендовано\b.+\b(подготовить|создать|разработать|оформить|согласовать|добавить|исправить|проверить|уточнить|сократить|сохранить|сделать|переформулировать)\b",
    r"\bтребуется\b.+\b(подготовить|создать|разработать|оформить|согласовать|добавить|исправить|проверить|уточнить|сократить|сделать|переформулировать)\b",
]

TASK_FALSE_POSITIVE = [
    r"\bзадача\s+(оптимизации|минимизации|максимизации|активизации|исследования|мст|сортировки)\b",
    r"\bглобальная\s+задача\b",
    r"\bпостановк\w*\s+задач\b",
    r"\bформулиров\w*\s+задач\b",
    r"\bматематическ\w*\s+задач\b",
    r"\bпохож\w*\s+задач\w*\b",
    r"\bэта\s+задача\b",
    r"\bинтересная\s+задача\b",
    r"\bкак\s+конкретно\s+задачу\b",
    r"\bя\s+думаю\b",
    r"\bмне\s+кажется\b",
    r"\bя\s+бы\s+(хотел|не\s+хотел)\b",
    r"\bя\s+хочу\b",
    r"\bя\s+решил\b",
    r"\bя\s+начал\b",
    r"\bя\s+попытался\b",
    r"\bне\s+знаю,?\s+как\b",
    r"\bнадо\s+подумать\b",
    r"\bнужно\s+просто\b",
    r"\bнадо\s+как-то\b",
    r"\bнужно\s+как-то\b",
    r"\bможно\s+(как-то|попробовать|оценить|описать|пояснить)\b",
    r"\bминимизировать\b",
    r"\bмаксимизировать\b",
    r"\bоптимальн\w*\b",
    r"\bдиаметр\w*\b",
    r"\bграф\w*\b",
    r"\bметрик\w*\b",
    r"\bмодель\w*\b",
    r"\bэвристик\w*\b",
    r"\bалгоритм\w*\b",
]

ANSWER_STARTERS = [
    r"^да\b", r"^нет\b", r"^хорошо\b", r"^ладно\b", r"^ок(ей)?\b",
    r"^понял\b", r"^поняла\b", r"^согласен\b", r"^согласна\b",
    r"^верно\b", r"^именно\b", r"^конечно\b", r"^готово\b",
    r"^принято\b", r"^понятно\b",
]

ANSWER_STATUS = [
    r"\bя\s+(сделаю|подготовлю|отправлю|проверю|исправлю|обновлю|настрою|доделаю|согласую|напишу)\b",
    r"\bмы\s+(сделаем|подготовим|отправим|проверим|исправим|обновим)\b",
    r"\b(сделано|готово|выполнено|завершено|загружено|отправлено|реализовано)\b",
    r"\bза\s+это\s+отвечает\b",
    r"\bуже\s+(готов|готова|сделал|сделала|отправил|отправила|выполнил|выполнила)\b",
    r"\bв\s+процессе\b",
    r"\bработаю\s+над\b",
    r"\bзанимаюсь\b",
]

PROGRESS_PATTERNS = [
    r"\bесть\s+умерен(ный|ное|ная)\s+прогресс\b",
    r"\bмы\s+уже\s+синхронизировались\b",
    r"\bсейчас\s+коротко\s+расскажу\b",
    r"\bна\s+прошлой\s+неделе\s+мы\s+обсуждали\b",
    r"\bслово\s+передаю\b",
    r"\bобновление\s+по\s+своему\s+блоку\b",
    r"\bпрогресс\s+хороший\b",
]

DEADLINE_PATTERNS = [
    r"\bдо\s+\d{1,2}[:.]\d{2}\b",
    r"\bдо\s+\d{1,2}[./]\d{1,2}(?:[./]\d{2,4})?\b",
    r"\bдо\s+(завтра|вечера|утра|обеда|понедельника|вторника|среды|четверга|пятницы|конца\s+дня|конца\s+недели|следующей\s+встречи)\b",
    r"\bк\s+(утру|вечеру|созвону|завтрашнему\s+созвону|следующей\s+встрече)\b",
    r"\b(сегодня|завтра|на\s+этой\s+неделе|после\s+обеда)\b",
]

SKIP_PATTERNS = [
    r"^\s*$",
    r"^включили\s+запис",
    r"^выключили\s+запис",
    r"^\s*дата\s*:",
    r"^\s*участники\s*:",
    r"^\s*повестка\s*:",
    r"^\s*тема\s*:",
    r"^\s*место\s*:",
]

TIMECODE_RE = re.compile(r"^\s*\[\d{2}:\d{2}:\d{2}\]\s*")
SPEAKER_PREFIX_RE = re.compile(r"^\s*(Говорящий|Спикер|Speaker)\s*\d*\s*:\s*", re.I)
PERSON_PREFIX_RE = re.compile(r"^\s*([А-ЯЁ][а-яё]+)\s*,\s*(.+)$")

FULL_SPEAKER_RE = re.compile(
    r"^\s*[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}(?:\s*\(голос\s*\d+\))?:\s*$"
)

HEADER_RE = re.compile(
    r"^\s*(протокол встречи|ход встречи|ход проведения встречи|повестка|"
    r"на встрече|рассмотренные вопросы|организация встречи|формат:|дата:|"
    r"встреча проходила|участники встречи|присутствовали|результаты встречи|"
    r"итог встречи|тема\s+\d+|следующая встреча|время:|"
    r"поставленные задачи|задачи к следующей встрече|что нужно сделать до следующей встречи|"
    r"в рамках продолжения работ определены следующие задачи)\b",
    re.I,
)

BAD_SHORT_QUESTIONS = {
    "да?",
    "нет?",
    "что?",
    "где?",
    "когда?",
    "зачем?",
    "почему?",
    "и все?",
    "чего?",
    "как?",
    "а?",
    "ну?",
    "правильно?",
    "да.",
    "нет.",
}

SUMMARY_OTHER_STARTS = (
    "обсуждение",
    "демонстрация",
    "рассмотрение",
    "на встрече",
    "в ходе встречи",
    "итог встречи",
    "результаты встречи",
    "отмечено",
    "требования к",
    "контур",
    "формализация",
    "математическая модель",
    "общее содержание",
    "название главы",
    "вводный абзац",
    "по итогам обсуждения",
    "по итогам показа",
    "текущая тема",
    "ход встречи",
    "присутствовали",
    "участники встречи",
    "дата следующей встречи",
    "следующая встреча",
)


def normalize(text: str) -> str:
    text = str(text).replace("\xa0", " ")
    text = TIMECODE_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip(" -–—")


def remove_only_speaker_prefix(text: str) -> str:
    return SPEAKER_PREFIX_RE.sub("", text).strip()


def is_full_speaker_line(text: str) -> bool:
    return FULL_SPEAKER_RE.match(text.strip()) is not None


def is_header_line(text: str) -> bool:
    return HEADER_RE.match(text.strip()) is not None


def is_bad_short_question(text: str) -> bool:
    return text.lower().strip() in BAD_SHORT_QUESTIONS


def is_summary_other(text: str) -> bool:
    lower = text.lower().strip()
    return any(lower.startswith(s) for s in SUMMARY_OTHER_STARTS)


def split_into_sentences(text: str) -> list[str]:
    text = normalize(text)
    if not text:
        return []

    parts = re.split(r"(?<=[.!?])\s+|\s*;\s*", text)
    result: list[str] = []

    for part in parts:
        part = normalize(part)

        if not part:
            continue

        if len(part) < MIN_LEN:
            continue

        if len(part.split()) < MIN_WORDS:
            continue

        result.append(part)

    return result


def is_filler_heavy(text: str) -> bool:
    words = [w.strip(".,!?;:()[]\"'«»").lower() for w in text.split()]
    words = [w for w in words if w]

    if len(words) <= 3:
        return False

    filler_count = sum(1 for w in words if w in FILLERS)
    return filler_count / max(len(words), 1) >= 0.45


def has_deadline(text: str) -> bool:
    lower = text.lower()
    return any(re.search(p, lower) for p in DEADLINE_PATTERNS)


def has_person_prefix(text: str) -> bool:
    return PERSON_PREFIX_RE.match(text) is not None


def has_imperative(text: str) -> bool:
    lower = text.lower()
    return any(re.search(rf"\b{re.escape(v)}\b", lower) for v in IMPERATIVE_WORDS)


def has_action_infinitive(text: str) -> bool:
    lower = text.lower()
    return any(re.search(rf"\b{re.escape(v)}\b", lower) for v in ACTION_INF)


def matches_protocol_task(text: str) -> bool:
    lower = text.lower()
    return any(re.search(p, lower) for p in PROTOCOL_TASK_PATTERNS)


def is_question(text: str) -> bool:
    lower = text.lower().strip()

    if is_bad_short_question(text):
        return False

    if text.endswith("?"):
        if len(text.split()) <= 2:
            return False
        return True

    return any(lower.startswith(w + " ") for w in QUESTION_WORDS)


def is_answer(text: str) -> bool:
    lower = text.lower().strip()

    if is_header_line(text):
        return False

    if is_summary_other(text):
        return False

    if lower.startswith("итог встречи"):
        return False

    if lower.startswith("обсуждение"):
        return False

    if lower.startswith("демонстрация"):
        return False

    if lower.startswith("рассмотрение"):
        return False

    if any(re.search(p, lower) for p in ANSWER_STARTERS):
        return True

    if any(re.search(p, lower) for p in ANSWER_STATUS):
        return True

    if lower.startswith("это "):
        return True

    if lower.startswith("вот,"):
        return True

    return False


def is_progress_or_other(text: str) -> bool:
    lower = text.lower()
    return any(re.search(p, lower) for p in PROGRESS_PATTERNS)


def is_false_task_context(text: str) -> bool:
    lower = text.lower()
    return any(re.search(p, lower) for p in TASK_FALSE_POSITIVE)


def is_task(text: str) -> bool:
    lower = text.lower().strip()

    if is_header_line(text):
        return False

    if is_summary_other(text):
        if lower.startswith("отмечено, что необходимо") or lower.startswith("отмечено, что нужно"):
            pass
        elif lower.startswith("рекомендовано"):
            pass
        else:
            return False

    if len(text.split()) <= 2:
        return False

    if is_question(text) or is_answer(text) or is_progress_or_other(text):
        return False

    if is_false_task_context(text):
        return False

    if is_filler_heavy(text):
        return False

    person = has_person_prefix(text)
    deadline = has_deadline(text)
    imperative = has_imperative(text)
    inf_action = has_action_infinitive(text)
    protocol_task = matches_protocol_task(text)

    if person and (imperative or inf_action or deadline):
        return True

    if imperative and (deadline or len(text.split()) >= 3):
        return True

    if protocol_task:
        return True

    if inf_action and deadline:
        return True

    if re.search(r"\b(нужно|надо|необходимо|требуется|прошу)\b", lower) and inf_action:
        return True

    if re.search(r"\b(следует|планируется|принято решение|рекомендовано)\b", lower) and inf_action:
        return True

    if re.match(
        r"^\s*(подготовка|проверка|согласование|разработка|реализация|оформление|создание|доработка|тестирование)\b",
        lower,
    ):
        return True

    return False


def classify_sentence(text: str) -> str:
    text = normalize(remove_only_speaker_prefix(text))

    if not text:
        return "skip"

    if len(text) < MIN_LEN:
        return "skip"

    if is_full_speaker_line(text):
        return "skip"

    if is_header_line(text):
        return "other"

    for p in SKIP_PATTERNS:
        if re.search(p, text.lower()):
            return "skip"

    if is_question(text):
        return "question"

    if is_task(text):
        return "task"

    if is_answer(text):
        return "answer"

    return "other"


def read_docx(path: Path) -> list[str]:
    if not DOCX_AVAILABLE:
        print(f"python-docx не установлен, пропускаю {path.name}")
        return []

    try:
        doc = DocxDocument(path)
    except Exception as e:
        print(f" Не удалось открыть {path.name}: {e}")
        return []

    chunks: list[str] = []

    for p in doc.paragraphs:
        t = normalize(p.text)
        if t:
            chunks.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(normalize(cell.text) for cell in row.cells if normalize(cell.text))
            if row_text:
                chunks.append(row_text)

    return chunks


def read_txt(path: Path) -> list[str]:
    try:
        return path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig").splitlines()


def read_file(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    return read_txt(path)


def read_jsonl(path: Path) -> list[str]:
    lines: list[str] = []

    try:
        raw_lines = path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        raw_lines = path.read_text(encoding="utf-8-sig").splitlines()

    for line in raw_lines:
        line = line.strip()

        if not line:
            continue

        try:
            item = json.loads(line)
        except Exception:
            continue

        text = normalize(str(item.get("text", "")))

        if text:
            lines.append(text)

    return lines


def process_line(raw: str, source: str) -> list[dict]:
    raw = normalize(raw)

    if not raw:
        return []

    raw = remove_only_speaker_prefix(raw)

    if is_full_speaker_line(raw):
        return []

    items: list[dict] = []

    for sent in split_into_sentences(raw):
        label = classify_sentence(sent)

        if label == "skip":
            continue

        items.append({
            "text": sent,
            "label": label,
            "source": source,
        })

    return items


def process_file(path: Path) -> list[dict]:
    lines = read_file(path)
    items: list[dict] = []

    for raw in lines:
        items.extend(process_line(raw, path.name))

    return items


def process_jsonl(path: Path) -> list[dict]:
    lines = read_jsonl(path)
    items: list[dict] = []

    for raw in lines:
        items.extend(process_line(raw, path.name))

    return items


def dedupe_items(items: list[dict]) -> list[dict]:
    seen = set()
    out = []

    for item in items:
        text = normalize(item["text"])
        key = text.lower()

        if not key or key in seen:
            continue

        seen.add(key)

        out.append({
            "text": text,
            "label": item["label"],
            "source": item.get("source", ""),
        })

    return out


def save_jsonl(items: Iterable[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8") as f:
        for item in items:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


def main() -> None:
    transcript_files = sorted([
        p for p in TRANSCRIPTS_DIR.iterdir()
        if p.suffix.lower() in {".txt", ".docx"}
    ])

    jsonl_files = [p for p in RAW_JSONL_FILES if p.exists()]

    if not transcript_files and not jsonl_files:
        print("Файлы не найдены.")
        print("Положите .txt/.docx в transcripts/ или .jsonl в корень проекта / datasets/")
        return

    all_items: list[dict] = []

    for path in transcript_files:
        items = process_file(path)
        print(f"Обработка: {path.name}... {len(items)} предложений")
        all_items.extend(items)

    for path in jsonl_files:
        items = process_jsonl(path)
        print(f"Обработка JSONL: {path.name}... {len(items)} предложений")
        all_items.extend(items)

    all_items = dedupe_items(all_items)
    save_jsonl(all_items, OUTPUT_PATH)

    stats: dict[str, int] = {}

    for item in all_items:
        stats[item["label"]] = stats.get(item["label"], 0) + 1

    print("\nГотово:")
    print(f"  {OUTPUT_PATH}")
    print("Распределение классов:")

    for label in ["task", "question", "answer", "other"]:
        print(f"  {label:8} {stats.get(label, 0)}")

    print("\nЧто дальше:")
    print("1. Откройте datasets/prelabeled.jsonl")
    print("2. Быстро проверьте task и answer")
    print("3. Запустите python balance_dataset.py")
    print("4. Потом python train_classifier.py")


if __name__ == "__main__":
    main()