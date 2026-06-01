import json
from pathlib import Path

from docx import Document

from scripts.export_annotation_queue import build_annotation_queue, save_jsonl


def test_annotation_queue_creates_pending_items(tmp_path: Path):
    input_dir = tmp_path / "transcripts"
    input_dir.mkdir()
    docx_path = input_dir / "meeting.docx"
    doc = Document()
    doc.add_paragraph("Да, я уже проверил эту часть.")
    doc.add_paragraph("За это отвечает Анна.")
    doc.save(docx_path)
    dataset = tmp_path / "pm_dataset.jsonl"
    dataset.write_text("", encoding="utf-8")

    queue = build_annotation_queue(input_dir, dataset)
    output = tmp_path / "annotation_queue.jsonl"
    save_jsonl(output, queue)

    assert output.exists()
    assert queue
    assert all(item["annotation_status"] == "pending" for item in queue)
    assert {item["suggested_label"] for item in queue} & {"answer", "responsible"}


def test_manual_seed_examples_are_valid_jsonl():
    rows = [
        json.loads(line)
        for line in Path("datasets/manual_seed_examples.jsonl").open(encoding="utf-8")
        if line.strip()
    ]

    assert len(rows) == 50
    assert {row["metadata"]["source"] for row in rows} == {"manual_seed"}
