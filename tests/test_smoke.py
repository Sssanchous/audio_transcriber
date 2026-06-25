import json
from io import BytesIO
import subprocess
import sys
from collections import Counter
from pathlib import Path

import joblib
from fastapi.testclient import TestClient

from pm_insights.api.main import app


def write_jsonl(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8") as fh:
        for row in rows:
            fh.write(json.dumps(row, ensure_ascii=False) + "\n")


def test_health_smoke():
    client = TestClient(app)
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["status"] == "ok"


def test_manual_examples_are_four_class_dataset():
    path = Path("datasets/sources/manual_examples.jsonl")
    rows = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
    labels = {row["label"] for row in rows}
    counts = Counter(row["label"] for row in rows)
    assert labels == {"task", "question", "answer", "other"}
    assert all(counts[label] >= 500 for label in labels)
    assert not any("????" in row["text"] for row in rows)
    assert not any("Рђ" in row["text"] or "Р°" in row["text"] for row in rows)


def test_training_dataset_is_four_class_dataset():
    path = Path("datasets/training_dataset.jsonl")
    rows = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
    labels = {row["label"] for row in rows}
    assert labels == {"task", "question", "answer", "other"}
    assert len(rows) > 2000
    assert not any("????" in row["text"] for row in rows)


def test_real_hard_examples_are_clean_four_class_examples():
    path = Path("datasets/sources/real_hard_examples.jsonl")
    rows = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
    labels = {row["label"] for row in rows}
    assert rows
    assert labels <= {"task", "question", "answer", "other"}
    assert "other" in labels
    assert not any("????" in row["text"] for row in rows)
    assert not any("Рђ" in row["text"] or "Р°" in row["text"] for row in rows)


def test_saved_baseline_predicts_real_meeting_phrases():
    model_path = Path("models/baseline_classifier/model.joblib")
    assert model_path.exists()
    model = joblib.load(model_path)
    checks = [
        ("Алексей, Анна, подготовь финальный отчет по аналитике до пятницы", "task"),
        ("Алексей, когда мы отправим презентацию клиенту?", "question"),
        ("Анна, я отправлю презентацию сегодня после обеда", "answer"),
        ("Коллеги, всем доброе утро. Начинаем еженедельную встречу по проекту.", "other"),
        ("Иван. Нет, релиз переносить не нужно", "answer"),
        ("Алексей, можно ли согласовать бюджет сегодня?", "question"),
    ]
    predictions = model.predict([text for text, _ in checks])
    assert list(predictions) == [expected for _, expected in checks]


def test_baseline_training_smoke(tmp_path):
    rows = []
    examples = {
        "task": ["Анна, подготовь отчёт.", "Иван, исправь ошибку.", "Мария, проверь сборку."],
        "question": ["Когда готов отчёт?", "Кто отвечает за сборку?", "Можно ли согласовать бюджет?"],
        "answer": ["Да, я подготовлю отчёт.", "Нет, переносить не нужно.", "Файл уже лежит в папке."],
        "other": ["Коллеги, доброе утро.", "Начинаем встречу.", "Сегодня обсуждаем статус."],
    }
    for label, texts in examples.items():
        for index, text in enumerate(texts):
            rows.append({"id": f"{label}_{index}", "text": text, "label": label})

    train = tmp_path / "train.jsonl"
    val = tmp_path / "val.jsonl"
    test = tmp_path / "test.jsonl"
    write_jsonl(train, rows)
    write_jsonl(val, rows)
    write_jsonl(test, rows)
    output = tmp_path / "model"

    result = subprocess.run(
        [
            sys.executable,
            "scripts/train_baseline_classifier.py",
            "--train",
            str(train),
            "--val",
            str(val),
            "--test",
            str(test),
            "--output",
            str(output),
        ],
        check=False,
        text=True,
        capture_output=True,
    )
    assert result.returncode == 0, result.stderr
    assert (output / "model.joblib").exists()
    assert (output / "metrics.json").exists()


def test_evaluation_scripts_skip_missing_references():
    scripts = [
        ("scripts/evaluate_asr_wer.py", "ASR evaluation skipped"),
        ("scripts/evaluate_task_extraction.py", "Task extraction evaluation skipped"),
        ("scripts/evaluate_sentiment.py", "Sentiment evaluation skipped"),
    ]
    for script, expected in scripts:
        result = subprocess.run(
            [
                sys.executable,
                script,
                "--reference",
                "eval_data/missing_reference.json",
                "--prediction",
                "results/example_result_fixed.json",
            ],
            check=False,
            text=True,
            capture_output=True,
        )
        assert result.returncode == 0
        assert expected in result.stdout


def test_optional_sentiment_and_topic_fallbacks(monkeypatch):
    from pm_insights import settings
    from pm_insights.nlp.sentiment import analyze_sentiment
    from pm_insights.nlp.topic_modeling import build_topics, extract_topics

    fragments = [{"fragment_index": 1, "text": "Отлично, все получилось по серверу."}]
    assert analyze_sentiment(fragments)[0]["sentiment"] in {"positive", "neutral", "negative"}

    monkeypatch.setattr(settings, "ENABLE_MODEL_FALLBACK", True)
    rubert_result = analyze_sentiment(fragments, engine="rubert")[0]
    assert rubert_result["engine"] in {"rubert", "rule_based"}

    fallback_topics = build_topics(fragments, engine="fallback")
    bertopic_topics = build_topics(fragments, engine="bertopic")
    assert isinstance(fallback_topics, list)
    assert isinstance(bertopic_topics, list)

    topic_result = extract_topics(fragments, engine="auto", max_topics=4)
    assert topic_result["source"] in {"bertopic", "embedding_clustering", "rule_based_fallback"}
    assert isinstance(topic_result["topics"], list)
    assert isinstance(topic_result["aspect_frequencies"], dict)
    assert isinstance(topic_result["warnings"], list)


def test_bertopic_allow_fit_and_save_model_route(monkeypatch):
    from pm_insights import settings
    from pm_insights.nlp import topic_modeling

    class FakeTopicPath:
        def exists(self):
            return True

        def __str__(self):
            return "models/topic_model"

    monkeypatch.setattr(settings, "TOPIC_MIN_FRAGMENTS", 1)
    monkeypatch.setattr(topic_modeling.settings, "TOPIC_MODEL_PATH", FakeTopicPath())

    def fake_fit(valid, max_topics, allow_fit=False, save_model=False):
        assert valid
        assert allow_fit is True
        assert save_model is True
        return [
            {
                "topic_id": 0,
                "topic_name": "медицинские анализы",
                "keywords": ["медицинские", "анализы"],
                "count": len(valid),
                "fragment_ids": [item["_topic_fragment_id"] for item in valid],
                "confidence": 0.8,
                "texts": [item["text"] for item in valid],
                "source_text": " ".join(item["text"] for item in valid),
            }
        ]

    monkeypatch.setattr(topic_modeling, "_fit_bertopic_topics", fake_fit)
    result = topic_modeling.extract_topics(
        [
            {"fragment_index": 1, "text": "Обсудили медицинские анализы, клинический журнал и обучение персонала."},
            {"fragment_index": 2, "text": "Отдельно разобрали стерилизацию инструментов и контроль качества."},
        ],
        engine="bertopic",
        allow_fit=True,
        save_model=True,
    )
    assert result["source"] == "bertopic"
    assert result["topics"]
    assert result.get("model_saved_to")


def test_bertopic_failure_warns_and_falls_back_to_embedding(monkeypatch):
    from pm_insights import settings
    from pm_insights.nlp import topic_modeling

    monkeypatch.setattr(settings, "TOPIC_MIN_FRAGMENTS", 1)

    def fake_fit(*args, **kwargs):
        raise RuntimeError("BERTopic optional dependencies are not installed.")

    def fake_embedding(valid, max_topics):
        return [
            {
                "topic_id": 0,
                "topic_name": "стерилизация инструментов",
                "keywords": ["стерилизация", "инструменты"],
                "count": len(valid),
                "fragment_ids": [item["_topic_fragment_id"] for item in valid],
                "confidence": 0.7,
                "texts": [item["text"] for item in valid],
                "source_text": " ".join(item["text"] for item in valid),
            }
        ]

    monkeypatch.setattr(topic_modeling, "_fit_bertopic_topics", fake_fit)
    monkeypatch.setattr(topic_modeling, "_embedding_cluster_topics", fake_embedding)
    result = topic_modeling.extract_topics(
        [
            {"fragment_index": 1, "text": "Обсудили стерилизацию инструментов и журнал контроля."},
            {"fragment_index": 2, "text": "Отдельно разобрали обучение персонала клиники."},
        ],
        engine="bertopic",
        allow_fit=True,
    )
    assert result["source"] == "embedding_clustering"
    assert any("BERTopic unavailable" in warning for warning in result["warnings"])


def test_unknown_domain_topic_fallback_is_not_empty():
    from pm_insights.nlp.topic_modeling import extract_topics

    fragments = [
        {"fragment_index": 1, "text": "На встрече обсуждали стерилизацию инструментов, журнал контроля и обучение персонала клиники."},
        {"fragment_index": 2, "text": "Отдельно разобрали график медицинских осмотров и подготовку кабинетов к проверке."},
        {"fragment_index": 3, "text": "Юридический блок включал согласование договора, персональные данные и хранение документов."},
    ]
    result = extract_topics(fragments, engine="rule_based", max_topics=3)
    assert result["source"] == "rule_based_fallback"
    assert result["topics"]
    assert result["aspect_frequencies"]


def test_topic_names_ignore_filler_words_and_keep_semantic_keywords():
    from pm_insights.nlp.topic_modeling import _topic_name_from_keywords

    texts = [
        "Именно ага раз нас сейчас интересуют параметры пласта, скин-фактор, дебит и проницаемость скважины.",
        "В модели нужно сравнить дебит, пластовое давление и параметры трещины для интерпретации данных.",
        "Ага, тогда смотрим интерпретацию данных, эталон и калькулированные значения по скважине.",
    ]
    name = _topic_name_from_keywords(["нас", "ага", "раз", "дебит", "скин"], texts)
    assert name != "нас ага раз"
    assert "ага" not in name
    assert "нас" not in name
    assert any(term in name for term in ["дебит", "скин"])


def test_commercial_phrase_topic_is_normalized_before_display():
    from pm_insights.nlp.topic_modeling import _finalize_topics

    topics = _finalize_topics(
        [
            {
                "topic_name": "Даты рабочих контракте",
                "keywords": ["рабочих", "контракте", "отгрузки", "поставки", "даты"],
                "count": 3,
                "fragment_ids": [1, 2],
                "confidence": 0.7,
                "source_text": "По контракту обсуждали рабочие даты отгрузки, поставки и платежей.",
            },
            {
                "topic_name": "даты платежей",
                "keywords": ["даты", "платежей", "оплаты"],
                "count": 2,
                "fragment_ids": [3],
                "confidence": 0.7,
                "source_text": "Нужно прописать точные даты платежей и оплаты.",
            },
        ],
        max_topics=5,
    )
    names = {topic["topic_name"].lower() for topic in topics}

    assert "даты рабочих контракте" not in names
    assert "сроки поставки и оплаты" in names
    assert "платежные условия" in names


def _topic_names_for(texts: list[str]) -> set[str]:
    from pm_insights.nlp.topic_modeling import extract_topics

    fragments = [{"fragment_index": index, "text": text} for index, text in enumerate(texts, start=1)]
    result = extract_topics(fragments, engine="rule_based", max_topics=8)
    return {topic["topic_name"].lower() for topic in result["topics"]}


def test_topic_names_are_meaningful_and_drop_filler_clusters():
    names = _topic_names_for(
        [
            "Именно получается вообще какие-то данные что-то, Илья сказал правило.",
            "Спрашиваю в теории и в принципе, можно отталкиваться от чего-нибудь.",
        ]
    )
    joined = " ".join(names)
    assert "именно получается вообще" not in joined
    assert "какие-то данных что-то" not in joined
    assert "илья что-нибудь" not in joined


def test_unknown_domain_topic_still_uses_meaningful_keywords():
    names = _topic_names_for(
        [
            "На встрече обсуждали стерилизацию инструментов, журнал контроля и обучение персонала клиники.",
            "Отдельно разобрали график медицинских осмотров и подготовку кабинетов к проверке.",
            "Юридический блок включал согласование договора, персональные данные и хранение документов.",
        ]
    )
    joined = " ".join(names)
    assert names
    assert "что-то" not in joined
    assert any(term in joined for term in ["обучение сотрудников", "договор", "документация", "стерилизац", "персонал"])


def test_project_technical_education_commercial_and_oil_gas_topics_are_clean():
    project = _topic_names_for(
        [
            "Подготовить отчет по релизу, исправить ошибку авторизации и проверить сервер.",
            "Обсудили бюджет, сроки, клиентские требования и релизную сборку.",
        ]
    )
    assert {"релиз", "бюджет", "сервер", "авторизация"} & project

    technical = _topic_names_for(
        [
            "Модель, параметры модели, точность, аппроксимация и эксперимент по данным.",
            "Нужно проверить устойчивость модели, выборку и расчетные параметры.",
        ]
    )
    assert {"параметры модели", "модель и аппроксимация", "точность модели", "данные и выборка"} & technical

    education = _topic_names_for(
        [
            "ВКР, постановка задачи, материалы и методы, глава и черновик документа.",
            "Обсуждали критерии качества, ограничения работы, защиту и требования комиссии.",
        ]
    )
    assert {"ВКР и документация".lower(), "постановка задачи", "материалы и методы", "критерии качества"} & education

    commercial = _topic_names_for(
        [
            "Обсудили договор, цену, платежные условия, предоплату и банковскую гарантию.",
            "Зафиксировали коммерческие условия, обязательства сторон и график поставки.",
        ]
    )
    assert {"платежные условия", "цена", "договор", "банковская гарантия"} & commercial

    oil_gas = _topic_names_for(
        [
            "Brent и Dated Brent, премия, дифференциал и ценовая формула по партии.",
            "Объемы поставки, тонн, отгрузка, терминал, фрахт, демередж и судно.",
            "Качество сырья, плотность, сера, инспекция и коносамент.",
        ]
    )
    assert {"объемы поставки", "ценовая формула", "фрахт и демередж", "качество сырья", "инспекция"} & oil_gas


def test_non_meeting_speech_topics_do_not_emit_filler_noise():
    names = _topic_names_for(
        [
            "Дорогие коллеги, поздравляю вас с праздником и благодарю за работу.",
            "Это информационное сообщение и короткое обращение к команде.",
        ]
    )
    joined = " ".join(names)
    assert "именно" not in joined
    assert "что-нибудь" not in joined
    assert names <= {"поздравление", "благодарность", "информационное сообщение", "обращение", "организационная информация"}


def test_topic_source_hidden_from_ui_and_exports():
    frontend = Path("frontend/src/pages/ResultPage.jsx").read_text(encoding="utf-8")
    report_builder = Path("src/pm_insights/export/report_builder.py").read_text(encoding="utf-8")
    forbidden = ["Источник тем", "topic_source", "rule_based_fallback", "embedding_clustering"]
    for text in (frontend, report_builder):
        for phrase in forbidden:
            assert phrase not in text


def test_fragment_classifier_rule_based_and_baseline_safe():
    from pm_insights.nlp.fragment_classifier import classify_fragment

    question = classify_fragment("Когда будет готов финальный отчет?", engine="rule_based")
    task = classify_fragment("Анна, подготовь отчет до пятницы.", engine="rule_based")
    assert question["label"] == "question"
    assert task["label"] == "task"

    baseline = classify_fragment("Коллеги, начинаем встречу по проекту.", engine="baseline")
    assert baseline["label"] in {"task", "question", "answer", "other"}


def test_pipeline_adds_processing_time(tmp_path):
    from pm_insights.meeting.pipeline import analyze_meeting

    audio = tmp_path / "meeting.mp3"
    audio.write_bytes(b"fake-audio")
    result = analyze_meeting(audio, transcript_text="Нужно подготовить отчет до пятницы.", output_dir=tmp_path)
    processing_time = result["metadata"]["processing_time"]
    assert "total_processing_seconds" in processing_time
    assert "estimated_1h_processing_minutes" in processing_time


def test_dashboard_project_summary_payload(monkeypatch):
    from pm_insights.api import main as api_main

    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(api_main.db, "list_meetings_with_results", lambda user_id=None: [])
    monkeypatch.setattr(api_main, "_DASHBOARD_CACHE", {})
    response = TestClient(app).get("/dashboard")
    assert response.status_code == 200
    payload = response.json()
    assert "kpi_targets" not in payload
    assert "kpi_status" not in payload
    assert "business_kpi" not in payload
    assert "projects" in payload
    assert "summary" in payload
    assert "task_trend" in payload
    assert "sentiment_trend" in payload
    assert "aspect_word_cloud" in payload


def test_json_export_endpoint(monkeypatch):
    from pm_insights.api import main as api_main

    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(api_main.db, "get_result", lambda meeting_id, user_id=None: {"meeting_id": meeting_id, "tasks": []})
    response = TestClient(app).get("/meetings/demo/export/json")
    assert response.status_code == 200
    assert response.json()["meeting_id"] == "demo"
    assert "attachment" in response.headers["content-disposition"]


def test_delete_meeting_endpoint_removes_owned_meeting(monkeypatch):
    from pm_insights.api import main as api_main

    deleted = []
    storage_deleted = []
    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(
        api_main.db,
        "get_meeting",
        lambda meeting_id, user_id=None: {"meeting_id": meeting_id, "stored_filename": "demo.mp3"} if meeting_id == "demo" else None,
    )
    monkeypatch.setattr(api_main.db, "delete_meeting", lambda meeting_id, user_id=None: deleted.append((meeting_id, user_id)) or True)
    monkeypatch.setattr(api_main.meeting_storage, "delete_meeting", lambda meeting_id, output_dir=None: storage_deleted.append(meeting_id) or True)

    response = TestClient(app).delete("/meetings/demo")

    assert response.status_code == 200
    assert response.json() == {"status": "deleted", "meeting_id": "demo"}
    assert deleted == [("demo", None)]
    assert storage_deleted == ["demo"]


def test_upload_queues_analysis_when_async_processing_enabled(monkeypatch):
    from pathlib import Path
    from pm_insights.api import main as api_main

    async def fake_save_upload(_file):
        return Path("uploads/demo.mp3"), ".mp3", 123

    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(api_main.settings, "ASYNC_PROCESSING", True)
    monkeypatch.setattr(api_main, "_save_upload", fake_save_upload)
    monkeypatch.setattr(
        api_main.db,
        "create_meeting",
        lambda **kwargs: {"meeting_id": kwargs["meeting_id"], "processing_status": "uploaded"},
    )
    monkeypatch.setattr(api_main.db, "log_processing", lambda *args, **kwargs: None)
    monkeypatch.setattr(api_main, "_enqueue_meeting_analysis", lambda meeting_id: "task-demo")

    response = TestClient(app).post(
        "/upload",
        data={"meeting_title": "Demo", "project_name": "Demo"},
        files={"file": ("demo.mp3", b"fake", "audio/mpeg")},
    )

    assert response.status_code == 200
    assert response.json()["status"] == "queued"
    assert response.json()["task_id"] == "task-demo"


def test_upload_keeps_sync_behavior_when_async_processing_disabled(monkeypatch):
    from pathlib import Path
    from pm_insights.api import main as api_main

    async def fake_save_upload(_file):
        return Path("uploads/demo.mp3"), ".mp3", 123

    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(api_main.settings, "ASYNC_PROCESSING", False)
    monkeypatch.setattr(api_main, "_save_upload", fake_save_upload)
    monkeypatch.setattr(
        api_main.db,
        "create_meeting",
        lambda **kwargs: {"meeting_id": kwargs["meeting_id"], "processing_status": "uploaded"},
    )
    monkeypatch.setattr(api_main.db, "log_processing", lambda *args, **kwargs: None)
    monkeypatch.setattr(api_main, "_enqueue_meeting_analysis", lambda meeting_id: (_ for _ in ()).throw(AssertionError("queue should not be used")))

    response = TestClient(app).post(
        "/upload",
        data={"meeting_title": "Demo", "project_name": "Demo"},
        files={"file": ("demo.mp3", b"fake", "audio/mpeg")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "uploaded"
    assert "task_id" not in payload


def test_status_endpoint_returns_current_processing_status(monkeypatch):
    from pm_insights.api import main as api_main

    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(
        api_main.db,
        "get_meeting",
        lambda meeting_id, user_id=None: {
            "meeting_id": meeting_id,
            "processing_status": "completed",
            "metadata": {"async_task_id": "task-demo"},
        },
    )

    response = TestClient(app).get("/meetings/demo/status")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "completed"
    assert payload["job_status"] == "success"
    assert payload["result_ready"] is True


def test_dynamic_repeated_aspects_uses_clean_quality_gate(monkeypatch):
    from pm_insights.api import main as api_main

    previous = {
        "meeting_id": "prev",
        "meeting_key": "series",
        "meeting_type": {"label": "education_consultation"},
        "metadata": {"meeting_info": {"meeting_key": "series"}},
        "metrics": {"tasks_count": 1, "questions_count": 1, "average_sentiment": 0.1, "topic_frequencies": {"Требования этой эту": 3}},
        "topics": [{"topic_name": "Диаметр интересно случайный", "keywords": ["диаметр", "граф", "топология"], "count": 1}],
        "aspects": [],
    }
    current = {
        "meeting_id": "current",
        "meeting_key": "series",
        "meeting_type": {"label": "education_consultation"},
        "metadata": {"meeting_info": {"meeting_key": "series"}},
        "metrics": {"tasks_count": 2, "questions_count": 1, "average_sentiment": 0.2, "topic_frequencies": {"Требования этой эту": 3}},
        "topics": [{"topic_name": "Диаметр интересно случайный", "keywords": ["диаметр", "граф", "топология"], "count": 1}],
        "aspects": [],
    }
    monkeypatch.setattr(api_main.db, "list_meetings_with_results", lambda user_id=None: [previous, current])
    monkeypatch.setattr(api_main, "_DYNAMIC_ANALYSIS_CACHE", {})
    monkeypatch.setattr(api_main, "_NORMALIZED_RESULT_CACHE", {})

    dynamic = api_main._dynamic_analysis_for_result(current, user_id=None)

    assert "Требования этой эту" not in dynamic["repeated_topics"]
    assert "диаметр и топология графа" in [topic.lower() for topic in dynamic["repeated_topics"]]


def test_dynamic_repeated_aspects_uses_clean_raspberry_topics(monkeypatch):
    from pm_insights.api import main as api_main

    base = {
        "meeting_key": "raspberry-series",
        "meeting_type": {"label": "education_consultation"},
        "metadata": {"meeting_info": {"meeting_key": "raspberry-series"}},
        "transcript": [
            {
                "fragment_index": 1,
                "text": "Обсуждали Raspberry Pi, настройку окружения, YOLO benchmark, камеры и FPS, QEMU и задачи для второкурсников.",
            }
        ],
        "metrics": {"average_sentiment": 0.1},
        "topics": [
            {"topic_name": "Скин-фактор и трещины", "keywords": ["скин-фактор", "трещины"], "count": 3},
            {"topic_name": "Параметры пласта", "keywords": ["параметры", "пласта"], "count": 3},
            {"topic_name": "пакете пакете", "keywords": ["пакете", "пакете"], "count": 3},
            {"topic_name": "пользоваться config raspberry", "keywords": ["config", "raspberry", "cpu"], "count": 3},
            {"topic_name": "камеры fps", "keywords": ["камеры", "fps"], "count": 3},
            {"topic_name": "qemu виртуализация", "keywords": ["qemu", "виртуализация"], "count": 3},
            {"topic_name": "второкурсниками много", "keywords": ["второкурсниками", "много"], "count": 3},
            {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 3},
        ],
        "aspects": [{"aspects": ["Параметры пласта", "пакете пакете", "второкурсниками много"]}],
    }
    previous = dict(base, meeting_id="prev")
    current = dict(base, meeting_id="current")
    monkeypatch.setattr(api_main.db, "list_meetings_with_results", lambda user_id=None: [previous, current])
    monkeypatch.setattr(api_main, "_DYNAMIC_ANALYSIS_CACHE", {})
    monkeypatch.setattr(api_main, "_NORMALIZED_RESULT_CACHE", {})

    dynamic = api_main._dynamic_analysis_for_result(current, user_id=None)
    repeated = {topic.lower() for topic in dynamic["repeated_topics"]}

    assert "скин-фактор и трещины" not in repeated
    assert "параметры пласта" not in repeated
    assert "пакете пакете" not in repeated
    assert "второкурсниками много" not in repeated
    assert "raspberry pi и настройка окружения" in repeated
    assert "yolo benchmark" in repeated
    assert "камеры и fps" in repeated
    assert "задачи для второкурсников" in repeated
    assert "qemu / виртуализация raspberry pi" in repeated


def test_report_builder_uses_clean_topics_not_stale_raw_sections():
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.export.report_builder import build_report_payload, export_docx, export_pdf, export_xlsx

    result = {
        "meeting_id": "raspberry-clean",
        "meeting_type": {"label": "education_consultation"},
        "metadata": {"meeting_info": {"meeting_title": "Raspberry ВКР", "project_name": "ВКР"}},
        "transcript": [
            {
                "fragment_index": 1,
                "text": "Обсуждали Raspberry Pi, YOLO benchmark, камеры, QEMU, OpenCV и задачи для второкурсников.",
            }
        ],
        "report_sections": [
            {
                "id": "aspects_topics",
                "title": "Аспекты и темы",
                "items": [
                    {"topic_name": "Скин-фактор и трещины", "count": 4},
                    {"topic_name": "пакете пакете", "count": 3},
                ],
            }
        ],
        "topics": [
            {"topic_name": "пользоваться config raspberry", "keywords": ["config", "raspberry", "cpu"], "count": 4},
            {"topic_name": "Скин-фактор и трещины", "keywords": ["скин-фактор", "трещины"], "count": 4},
            {"topic_name": "пакете пакете", "keywords": ["пакете", "пакете"], "count": 3},
            {"topic_name": "камеры fps", "keywords": ["камеры", "fps"], "count": 3},
            {"topic_name": "qemu виртуализация", "keywords": ["qemu", "виртуализация"], "count": 3},
            {"topic_name": "второкурсниками много", "keywords": ["второкурсниками", "много"], "count": 3},
            {"topic_name": "YOLO benchmark", "keywords": ["yolo", "benchmark"], "count": 5},
        ],
        "aspect_frequencies": {"Параметры пласта": 9, "пакете пакете": 5},
        "aspects": [{"aspects": ["Параметры пласта", "второкурсниками много", "камеры fps", "qemu виртуализация"]}],
    }

    report = build_report_payload(result)
    rendered = "\n".join(row.get("title", "") for row in report["compact"]["topics"])
    rendered += "\n" + "\n".join(row.get("aspect", "") for row in report["compact"]["aspects"])

    assert "Скин-фактор" not in rendered
    assert "Параметры пласта" not in rendered
    assert "пакете пакете" not in rendered
    assert "Raspberry Pi и настройка окружения" in rendered
    assert "YOLO benchmark" in rendered
    assert "камеры и FPS" in rendered
    assert "задачи для второкурсников" in rendered
    assert "QEMU / виртуализация Raspberry Pi" in rendered

    document = Document(BytesIO(export_docx(report)))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    docx_text += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    workbook = load_workbook(BytesIO(export_xlsx(report)))
    xlsx_text = "\n".join(
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )
    pdf_bytes = export_pdf(report)
    assert pdf_bytes.startswith(b"%PDF")
    for text in (docx_text, xlsx_text):
        assert "Скин-фактор и трещины" not in text
        assert "Параметры пласта" not in text
        assert "пакете пакете" not in text
        assert "второкурсниками много" not in text
        assert "Raspberry Pi и настройка окружения" in text
        assert "YOLO benchmark" in text
        assert "камеры и FPS" in text
        assert "задачи для второкурсников" in text
        assert "QEMU / виртуализация Raspberry Pi" in text
    for forbidden in ("Скин-фактор и трещины", "Параметры пласта", "пакете пакете", "второкурсниками много"):
        assert forbidden.encode("utf-8") not in pdf_bytes


def test_binary_export_endpoints(monkeypatch):
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.api import main as api_main

    result = {
        "meeting_id": "demo",
        "metadata": {"meeting_info": {"meeting_title": "Demo", "project_name": "PM Insights"}},
        "clean_tasks": [{"title": "Подготовить отчет", "summary": "Кратко оформить результаты"}],
        "report_sections": [{"id": "tasks", "title": "Задачи", "items": [{"title": "Подготовить отчет"}]}],
        "transcript": [{"start": 0, "end": 1, "text": "Нужно подготовить отчет."}],
    }
    monkeypatch.setattr(api_main.settings, "REQUIRE_AUTH", False)
    monkeypatch.setattr(api_main.db, "get_result", lambda meeting_id, user_id=None: result | {"meeting_id": meeting_id})
    client = TestClient(app)

    xlsx = client.get("/meetings/demo/export/xlsx")
    assert xlsx.status_code == 200
    workbook = load_workbook(BytesIO(xlsx.content))
    assert "Summary" in workbook.sheetnames
    assert "Tasks" in workbook.sheetnames
    assert workbook["Summary"]["A2"].value == "Система"

    docx = client.get("/meetings/demo/export/docx")
    assert docx.status_code == 200
    document = Document(BytesIO(docx.content))
    assert any("PM Insights" in paragraph.text for paragraph in document.paragraphs)
    assert not any("????" in paragraph.text for paragraph in document.paragraphs)

    pdf = client.get("/meetings/demo/export/pdf")
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF")


def test_export_recovers_corrupted_section_titles():
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.export.report_builder import (
        _section_title,
        build_report_payload,
        export_docx,
        export_pdf,
        export_xlsx,
        register_cyrillic_pdf_font,
    )

    damaged_sections = [
        {"id": "tasks", "title": "??????", "items": [{"title": "Проверить модель"}]},
        {"id": "questions_answers", "title": "??????? ? ??????", "items": [{"question_title": "Что проверить?", "answer_summary": "Частный случай."}]},
        {"id": "deadlines", "title": "???????? / ????????? ???????", "items": [{"deadline": "четверг 19:30"}]},
        {"id": "topics", "title": "???????? ????", "items": [{"topic_name": "модель"}]},
        {"id": "_".join(("review", "items")), "title": "??????? ????????", "items": [{"text": "спорный элемент"}]},
        {"id": "transcript", "title": "??????????", "items": []},
    ]
    assert _section_title(damaged_sections[0]) == "Задачи"
    assert _section_title(damaged_sections[1]) == "Вопросы и ответы"

    result = {
        "meeting_id": "technical-demo",
        "meeting_type": {"label": "technical_research"},
        "metadata": {"meeting_info": {"meeting_title": "Техническая встреча", "project_name": "ВКР"}},
        "report_sections": damaged_sections,
        "transcript": [{"start": 0, "end": 1, "text": "Проверить модель."}],
    }
    report = build_report_payload(result)
    docx_bytes = export_docx(report)
    document = Document(BytesIO(docx_bytes))
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for expected in [
        "Основные показатели",
        "Основные темы",
        "Задачи",
        "Ключевые вопросы и ответы",
        "Дедлайны",
    ]:
        assert expected in paragraphs
    assert "Транскрипт" not in paragraphs
    assert "????" not in paragraphs

    workbook = load_workbook(BytesIO(export_xlsx(report)))
    assert workbook["Summary"]["A2"].value == "Система"
    summary_values = [cell.value for row in workbook["Summary"].iter_rows() for cell in row if cell.value is not None]
    assert "ВКР" in summary_values

    assert register_cyrillic_pdf_font()
    pdf = export_pdf(report)
    assert pdf.startswith(b"%PDF")
    assert len(pdf) > 1000


def test_result_page_and_exports_do_not_render_short_summary_block():
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.export.report_builder import build_report_payload, export_docx, export_xlsx

    result = {
        "meeting_id": "summary-demo",
        "meeting_type": {"label": "project_meeting", "display_name": "проектная встреча"},
        "metadata": {"meeting_info": {"meeting_title": "Демо", "project_name": "PM Insights"}},
        "analysis_summary": {
            "meeting_type": "project_meeting",
            "main_topics": ["релиз", "сроки"],
            "summary_text": "Шаблонное краткое содержание не должно попадать в UI или экспорт.",
        },
        "clean_tasks": [{"title": "Подготовить отчет"}],
        "transcript": [{"start": 0, "end": 1, "text": "Нужно подготовить отчет."}],
    }
    frontend_source = Path("frontend/src/pages/ResultPage.jsx").read_text(encoding="utf-8")
    export_source = Path("src/pm_insights/export/report_builder.py").read_text(encoding="utf-8")
    assert "Краткое содержание" not in frontend_source
    assert "Краткое содержание" not in export_source

    report = build_report_payload(result)
    document = Document(BytesIO(export_docx(report)))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Краткое содержание" not in docx_text
    assert "Шаблонное краткое содержание" not in docx_text
    assert "Основные темы" in docx_text

    workbook = load_workbook(BytesIO(export_xlsx(report)))
    summary_values = [
        str(cell.value)
        for row in workbook["Summary"].iter_rows()
        for cell in row
        if cell.value is not None
    ]
    assert "Краткое содержание" not in summary_values
    assert "Тип" in summary_values


def test_result_page_and_exports_hide_review_debug_in_default_view():
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.export.report_builder import build_report_payload, export_docx, export_pdf, export_xlsx

    needs_review = "_".join(("needs", "review"))
    low_confidence_reason = "_".join(("low", "confidence", "or", "missing", "owner"))
    forbidden = [
        "Краткое содержание",
        "Требует проверки",
        "Нужна ручная проверка",
        "Проверка",
        "review_items",
        "needs_review",
        low_confidence_reason,
    ]
    result = {
        "meeting_id": "review-demo",
        "meeting_type": {"label": "project_meeting", "display_name": "проектная встреча"},
        "metadata": {"meeting_info": {"meeting_title": "Демо", "project_name": "PM Insights"}},
        "analysis_summary": {
            "meeting_type": "project_meeting",
            "main_topics": ["релиз"],
            "requires_manual_review": True,
        },
        "clean_tasks": [
            {
                "title": "Подготовить отчет",
                "summary": low_confidence_reason,
                "status": needs_review,
                needs_review: True,
                "_".join(("review", "required")): True,
                "reason": low_confidence_reason,
            }
        ],
        "report_sections": [
            {
                "id": "_".join(("review", "items")),
                "title": "Служебные элементы",
                "items": [{"text": low_confidence_reason}],
            }
        ],
        "transcript": [{"start": 0, "end": 1, "text": "Нужно подготовить отчет."}],
    }

    frontend_source = Path("frontend/src/pages/ResultPage.jsx").read_text(encoding="utf-8")
    export_source = Path("src/pm_insights/export/report_builder.py").read_text(encoding="utf-8")
    for text in (frontend_source, export_source):
        for phrase in forbidden:
            assert phrase not in text
    assert "Режим проверки" in frontend_source

    report = build_report_payload(result)
    document = Document(BytesIO(export_docx(report)))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for phrase in forbidden:
        assert phrase not in docx_text

    workbook = load_workbook(BytesIO(export_xlsx(report)))
    xlsx_text = "\n".join(
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )
    for phrase in forbidden:
        assert phrase not in xlsx_text

    pdf = export_pdf(report)
    assert pdf.startswith(b"%PDF")
    assert low_confidence_reason.encode("utf-8") not in pdf
    assert b"needs_review" not in pdf


def test_compact_exports_exclude_transcript_fragments_and_limit_sections():
    from docx import Document
    from openpyxl import load_workbook
    from pm_insights.export.report_builder import build_report_payload, export_docx, export_pdf, export_xlsx

    long_transcript = "Так, я бы хотел обсудить с вами пункт ВКР и полный текст расшифровки."
    oil_intro = "Доброе утро, коллеги! Давайте начинать"
    result = {
        "meeting_id": "compact-demo",
        "meeting_type": {"label": "education_consultation"},
        "metadata": {
            "meeting_info": {"meeting_title": "ВКР", "project_name": "ВКР", "meeting_date": "2026-06-02"},
            "duration_seconds": 3816.31,
            "asr_model": "large-v3",
            "processing_time": {"total_processing_seconds": 417.824},
        },
        "clean_tasks": [
            {"title": "Описать вариант в ВКР", "deadline": "четверг"},
            {"title": "Описать вариант в ВКР", "deadline": "пятница"},
            {"title": "Это необходимо сделать"},
            *[{"title": f"Действие {index}", "deadline": "четверг"} for index in range(12)],
        ],
        "clean_questions_answers": [
            {
                "question_title": "Что-то там минимизировать, что-то там максимизировать",
                "answer_summary": "Слабый фрагмент не должен попадать в отчёт.",
                "status": "answered",
            },
            {
                "question_title": "что у нас есть какие-то необходимые критерии",
                "answer_summary": "Слабый фрагмент не должен попадать в отчёт.",
                "status": "answered",
            },
            *[
                {
                "question_title": f"Как проверить постановку задачи {index}?",
                "answer_summary": "Проверить ограничения, критерии качества и структуру ВКР.",
                "status": "answered",
                }
                for index in range(12)
            ],
        ],
        "clean_deadlines": [
            {"deadline": f"срок {index}", "kind": "meeting_time", "context": "контекст " * 40}
            for index in range(12)
        ],
        "clean_topics": [
            {"topic_name": "Требования этой эту", "count": 5, "keywords": ["ВКР", "требования"]},
            {"topic_name": "Причем понял честно", "count": 4, "keywords": ["мусор"]},
            {"topic_name": "Диаметр интересно случайный", "count": 3, "keywords": ["д", "и", "а", "м", "е", "т", "р", ",", " ", "г", "р", "а", "ф"]},
            {"topic_name": "Только использую используют", "count": 3, "keywords": ["мусор"]},
            *[
                {"topic_name": f"полезная тема {index}", "count": index + 1, "keywords": ["модель", "данные"]}
                for index in range(12)
            ],
        ],
        "clean_aspects": [
            {"title": "требования к ВКР", "count": 5, "keywords": ["ВКР", "требования"]},
            {"title": "диаметр и топология графа", "count": 3, "keywords": ["диаметр", "граф"]},
        ],
        "clean_sentiment": [
            {"sentiment": "negative" if index % 5 == 0 else "neutral", "text": f"длинный фрагмент тональности {index} {long_transcript}"}
            for index in range(96)
        ],
        "analysis_summary": {"average_sentiment": -0.12},
        "transcript": [
            {"start": 0, "end": 1, "text": long_transcript},
            {"start": 1, "end": 2, "text": oil_intro},
        ],
    }

    report = build_report_payload(result)
    assert report["compact"]["meeting"]["duration"] == "1 ч 3 мин 36 сек"
    assert report["compact"]["meeting"]["processing_time"] == "6 мин 58 сек"
    assert len(report["compact"]["qa"]) == 6
    assert len(report["compact"]["topics"]) <= 8
    assert report["compact"]["aspects"] == []
    assert len(report["compact"]["deadlines"]) == 10
    assert [row["title"] for row in report["compact"]["tasks"]].count("Описать вариант в ВКР") == 1

    document = Document(BytesIO(export_docx(report)))
    docx_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    ) + "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    forbidden = [
        "Транскрипт",
        long_transcript,
        oil_intro,
        "Тональность: 96",
        "Аспекты и темы: 56",
        "Причем понял честно",
        "Только использую используют",
        "Что-то там минимизировать",
        "что у нас есть какие-то необходимые критерии",
        "[\"д\"",
        "Аспекты обсуждения",
    ]
    for phrase in forbidden:
        assert phrase not in docx_text
    for phrase in [
        "Основные показатели",
        "Основные темы",
        "Ключевые вопросы и ответы",
        "Тональность",
        "Позитив",
        "Нейтрально",
        "Негатив",
        "Средний тон",
        "диаметр и топология графа",
        "диаметр, граф",
    ]:
        assert phrase in docx_text

    workbook = load_workbook(BytesIO(export_xlsx(report)))
    assert "Transcript" not in workbook.sheetnames
    assert set(["Summary", "Tasks", "Questions_Answers", "Deadlines", "Topics", "Sentiment"]).issubset(workbook.sheetnames)
    assert workbook["Questions_Answers"].max_row <= 7
    assert workbook["Topics"].max_row <= 11
    assert workbook["Deadlines"].max_row <= 11
    xlsx_text = "\n".join(
        str(cell.value)
        for sheet in workbook.worksheets
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )
    for phrase in forbidden:
        assert phrase not in xlsx_text

    pdf = export_pdf(report)
    assert pdf.startswith(b"%PDF")
    assert "Транскрипт".encode("utf-8") not in pdf


def test_export_duration_formatting():
    from pm_insights.export.report_builder import format_duration

    assert format_duration(3816.31) == "1 ч 3 мин 36 сек"
    assert format_duration(417.824) == "6 мин 58 сек"


def test_export_keywords_are_not_rendered_as_character_arrays():
    from pm_insights.export.report_builder import format_keywords, normalize_keywords

    assert normalize_keywords("диаметр, граф, топология") == ["диаметр", "граф", "топология"]
    assert normalize_keywords([{"word": "диаметр"}, {"keyword": "граф"}]) == ["диаметр", "граф"]
    assert normalize_keywords(["д", "и", "а", "м", "е", "т", "р", ",", " ", "г", "р", "а", "ф"]) == ["диаметр", "граф"]
    formatted = format_keywords(["з", "н", "а", "ю", ",", " ", "м", "о", "д", "е", "л", "и"])
    assert "[\"з\"" not in formatted
    assert "знаю, модели" in formatted


def test_rubert_training_script_help():
    result = subprocess.run(
        [sys.executable, "scripts/train_rubert_classifier.py", "--help"],
        check=False,
        text=True,
        capture_output=True,
    )
    assert result.returncode == 0
    assert "Train RuBERT tiny classifier" in result.stdout


def test_prepare_training_maps_protocol_discussion_to_other():
    from scripts.prepare_training_dataset import normalize_record

    row = {"text": "Обсуждение результатов внедрения Qdrant", "label": "discussion_item"}
    normalized = normalize_record(row, 1, "protocol_dataset")
    assert normalized["label"] == "other"
    assert normalized["original_label"] == "discussion_item"
