from pathlib import Path

from docx import Document

from pm_insights.dataset.reader import read_docx


def test_docx_reader_reads_paragraphs_and_tables(tmp_path: Path):
    path = tmp_path / "meeting.docx"
    doc = Document()
    doc.add_paragraph("Нужно подготовить отчёт до пятницы.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Вопрос"
    table.cell(0, 1).text = "Когда будет готов макет?"
    doc.save(path)

    blocks = read_docx(path)

    assert any("подготовить отчёт" in block.text for block in blocks)
    assert any("Когда будет готов макет" in block.text for block in blocks)
