import json
from pathlib import Path

from docx import Document

from pm_insights.dataset.builder import build_dataset


def test_dataset_builder_creates_valid_jsonl(tmp_path: Path):
    input_dir = tmp_path / "transcripts"
    input_dir.mkdir()
    path = input_dir / "meeting.docx"
    doc = Document()
    doc.add_paragraph("Нужно подготовить отчёт до пятницы.")
    doc.add_paragraph("Протокол встречи")
    doc.add_paragraph("Когда будет готов макет?")
    doc.save(path)
    output = tmp_path / "dataset.jsonl"

    records, stats = build_dataset(input_dir, output, min_length=5)

    assert output.exists()
    assert len(records) == 2
    assert stats["labels"]["deadline"] == 1
    rows = [json.loads(line) for line in output.read_text(encoding="utf-8").splitlines()]
    assert all(row["id"] and row["text"] and row["label"] for row in rows)
