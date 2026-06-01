from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - covered by dependency-missing environments
    Document = None


@dataclass(frozen=True)
class RawBlock:
    source_file: str
    paragraph_index: int
    text: str
    block_type: str = "paragraph"


def find_docx_files(input_dir: str | Path) -> list[Path]:
    root = Path(input_dir)
    if not root.exists():
        return []
    return sorted(p for p in root.rglob("*.docx") if not p.name.startswith("~$"))


def read_docx(path: str | Path) -> list[RawBlock]:
    if Document is None:
        raise RuntimeError("python-docx is required to read .docx files")

    docx_path = Path(path)
    document = Document(docx_path)
    blocks: list[RawBlock] = []
    index = 0

    for paragraph in document.paragraphs:
        index += 1
        text = paragraph.text.strip()
        if text:
            blocks.append(RawBlock(docx_path.name, index, text, "paragraph"))

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                index += 1
                blocks.append(
                    RawBlock(
                        docx_path.name,
                        index,
                        " | ".join(cells),
                        f"table:{table_index}:{row_index}",
                    )
                )

    return blocks


def read_docx_folder(input_dir: str | Path) -> tuple[list[RawBlock], list[dict]]:
    blocks: list[RawBlock] = []
    errors: list[dict] = []

    for path in find_docx_files(input_dir):
        try:
            blocks.extend(read_docx(path))
        except Exception as exc:
            errors.append({"source_file": path.name, "error": str(exc)})

    return blocks, errors
