from __future__ import annotations

import json
import re
from pathlib import Path

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


ROOT_DIR = Path(".")
TRANSCRIPTS_DIR = ROOT_DIR / "transcripts"
DATASETS_DIR = ROOT_DIR / "datasets"

INPUT_JSONL_FILES = [
    ROOT_DIR / "dataset_from_protocols.jsonl",
    DATASETS_DIR / "raw.jsonl",
    DATASETS_DIR / "prelabeled_raw.jsonl",
    DATASETS_DIR / "manual_tasks.jsonl",
]

OUTPUT_PATH = DATASETS_DIR / "full_clean.jsonl"

VALID_LABELS = {"task", "question", "answer", "other"}

MIN_TEXT_LEN = 6
MIN_WORDS = 2

TIMECODE_RE = re.compile(r"^\s*\[\d{2}:\d{2}:\d{2}\]\s*")
SPEAKER_PREFIX_RE = re.compile(r"^\s*(Говорящий|Спикер|Speaker)\s*\d*\s*:\s*", re.I)

FULL_SPEAKER_RE = re.compile(
    r"^\s*[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}(?:\s*\(голос\s*\d+\))?:\s*$"
)

HEADER_RE = re.compile(
    r"^\s*(протокол|протокол встречи|рабочая встреча|ход встречи|повестка|"
    r"дата|участники|цель встречи|тема|место|формат|итог встречи|"
    r"результаты встречи|поставленные задачи|следующая встреча)\b",
    re.I,
)

QUESTION_WORDS = (
    "кто", "что", "где", "когда", "почему", "зачем", "как",
    "какой", "какая", "какое", "какие", "сколько",
    "можно ли", "нужно ли", "успеем ли", "надо ли",
    "в чём", "в чем", "что за", "куда", "откуда",
)

TASK_STARTS = (
    "подготовить",
    "доработать",
    "исправить",
    "проверить",
    "добавить",
    "реализовать",
    "сделать",
    "согласовать",
    "переписать",
    "оформить",
    "протестировать",
    "посмотреть",
    "уточнить",
    "сформировать",
    "описать",
    "разработать",
    "отобразить",
    "сохранить",
    "сократить",
    "переформулировать",
    "переделать",
    "подписать",
    "найти",
    "дополнить",
    "задокументировать",
    "написать",
    "провести",
    "изучить",
    "назначить",
    "разместить",
    "получить",
    "запросить",
    "забронировать",
)

TASK_MARKERS = (
    "нужно",
    "надо",
    "необходимо",
    "требуется",
    "следует",
    "рекомендовано",
    "принято решение",
    "к следующей встрече",
    "до следующей встречи",
)

ANSWER_PATTERNS = [
    r"^\s*(да|нет|хорошо|ладно|понял|поняла|понятно|ок|окей)\b",
    r"^\s*(конечно|верно|именно|согласен|согласна|принято)\b",
    r"\bя\s+(сделаю|подготовлю|отправлю|проверю|исправлю|обновлю|напишу|посмотрю|уточню|доделаю|согласую)\b",
    r"\bмы\s+(сделаем|подготовим|отправим|проверим|обновим|исправим)\b",
    r"\bуже\s+(сделал|сделала|готов|готова|отправил|отправила|проверил|проверила)\b",
    r"\b(сделано|готово|выполнено|завершено|загружено|отправлено|реализовано)\b",
    r"\bза\s+это\s+отвечает\b",
    r"\b(возьму|беру)\s+.*\bна\s+себя\b",
]

BAD_SHORT = {
    "да",
    "нет",
    "ок",
    "окей",
    "ага",
    "угу",
    "так",
    "ну",
    "вот",
    "ладно",
    "хорошо",
    "понятно",
}


def normalize_text(text: str) -> str:
    text = str(text).replace("\xa0", " ")
    text = TIMECODE_RE.sub("", text)
    text = SPEAKER_PREFIX_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip(" -–—")


def is_bad_text(text: str) -> bool:
    lower = text.lower().strip(" .,!?:;—-")

    if not text:
        return True

    if len(text) < MIN_TEXT_LEN:
        return True

    if len(text.split()) < MIN_WORDS:
        return True

    if lower in BAD_SHORT:
        return True

    if FULL_SPEAKER_RE.match(text):
        return True

    return False


def is_header(text: str) -> bool:
    return HEADER_RE.match(text.strip()) is not None


def split_sentences(text: str) -> list[str]:
    text = normalize_text(text)

    if not text:
        return []

    parts = re.split(r"(?<=[.!?])\s+|\s*;\s*", text)
    result = []

    for part in parts:
        part = normalize_text(part)

        if is_bad_text(part):
            continue

        result.append(part)

    return result


def looks_question(text: str) -> bool:
    text = normalize_text(text)
    lower = text.lower()

    if text.endswith("?"):
        return len(text.split()) > 2

    return any(lower.startswith(w + " ") for w in QUESTION_WORDS)


def looks_task(text: str) -> bool:
    text = normalize_text(text)
    lower = text.lower().strip()

    if looks_question(text):
        return False

    if lower.startswith(TASK_STARTS):
        return True

    if any(marker in lower for marker in TASK_MARKERS):
        if any(action in lower for action in TASK_STARTS):
            return True

    return False


def looks_answer(text: str) -> bool:
    text = normalize_text(text)
    lower = text.lower()

    if looks_question(text):
        return False

    if looks_task(text):
        return False

    return any(re.search(pattern, lower) for pattern in ANSWER_PATTERNS)


def auto_label(text: str, fallback_label: str = "other") -> str:
    text = normalize_text(text)

    if is_header(text):
        return "other"

    if looks_question(text):
        return "question"

    if looks_task(text):
        return "task"

    if looks_answer(text):
        return "answer"

    if fallback_label in VALID_LABELS:
        return fallback_label

    return "other"


def read_jsonl(path: Path) -> list[dict]:
    items = []

    if not path.exists():
        return items

    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()

            if not line:
                continue

            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue

            text = normalize_text(item.get("text", ""))
            label = str(item.get("label", "other")).strip()

            if label not in VALID_LABELS:
                label = "other"

            if is_bad_text(text):
                continue

            items.append({
                "text": text,
                "label": label,
                "source": path.name,
            })

    return items


def read_txt(path: Path) -> list[str]:
    try:
        return path.read_text(encoding="utf-8").splitlines()
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig").splitlines()


def read_docx(path: Path) -> list[str]:
    if not DOCX_AVAILABLE:
        print(f"python-docx не установлен, пропускаю {path.name}")
        return []

    try:
        doc = DocxDocument(path)
    except Exception as e:
        print(f"Не удалось открыть {path.name}: {e}")
        return []

    chunks = []

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(
                normalize_text(cell.text)
                for cell in row.cells
                if normalize_text(cell.text)
            )
            if row_text:
                chunks.append(row_text)

    return chunks


def read_transcript_file(path: Path) -> list[str]:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt(path)

    if suffix == ".docx":
        return read_docx(path)

    return []


def process_transcripts() -> list[dict]:
    items = []

    if not TRANSCRIPTS_DIR.exists():
        return items

    transcript_files = sorted(
        p for p in TRANSCRIPTS_DIR.iterdir()
        if p.suffix.lower() in {".txt", ".docx"}
    )

    for path in transcript_files:
        lines = read_transcript_file(path)
        count_before = len(items)

        for line in lines:
            for sent in split_sentences(line):
                label = auto_label(sent, "other")
                items.append({
                    "text": sent,
                    "label": label,
                    "source": path.name,
                })

        print(f"transcripts: {path.name} -> {len(items) - count_before}")

    return items


def dedupe_items(items: list[dict]) -> list[dict]:
    seen = set()
    result = []

    for item in items:
        text = normalize_text(item.get("text", ""))
        label = item.get("label", "other")
        source = item.get("source", "")

        if is_bad_text(text):
            continue

        key = text.lower()

        if key in seen:
            continue

        seen.add(key)

        result.append({
            "text": text,
            "label": label,
            "source": source,
        })

    return result


def save_jsonl(items: list[dict], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    with path.open("w", encoding="utf-8") as f:
        for item in items:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


def main() -> None:
    DATASETS_DIR.mkdir(exist_ok=True)

    all_items = []

    for path in INPUT_JSONL_FILES:
        loaded = read_jsonl(path)
        print(f"jsonl: {path} -> {len(loaded)}")
        all_items.extend(loaded)

    transcript_items = process_transcripts()
    all_items.extend(transcript_items)

    all_items = dedupe_items(all_items)

    save_jsonl(all_items, OUTPUT_PATH)

    stats = {}

    for item in all_items:
        stats[item["label"]] = stats.get(item["label"], 0) + 1

    print("\nГотово.")
    print(f"Файл: {OUTPUT_PATH}")
    print(f"Всего записей: {len(all_items)}")
    print("Распределение:")

    for label in ["task", "question", "answer", "other"]:
        print(f"  {label}: {stats.get(label, 0)}")


if __name__ == "__main__":
    main()